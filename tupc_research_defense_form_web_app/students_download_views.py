from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required
from django.contrib.auth.decorators import user_passes_test
from django.core.mail import send_mail
from django.conf import settings

from .forms import *
from .models import *

# Import Date & Time
from datetime import date

# Import Docx and PDF Convert
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from docx.shared import Inches
import qrcode
import os
import subprocess
import base64
from django.core.files.storage import FileSystemStorage
import cv2

today = date.today()
date_today = today.strftime("%B %d, %Y")

# Student - BET3 - Topic Defense - Panel Invitation - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentDownloadPanelInvitationBet3(request, id):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
        get_research_titles = ResearchTitle.objects.all().filter(student_leader_username=current_user.username)
        get_panel_invitation = TitlePanelInvitation.objects.get(id=int(id), student_leader_username=current_user.username)

    except:
        return redirect("student-panel-invitation-bet3")

    ############## BET-3 PANEL INVITATION DATA ##############
    date_submitted = get_panel_invitation.form_date_sent

    student_member_list = [get_panel_invitation.student_leader_full_name]
    course = get_student_leader_data.course.replace("Engineering", "Eng.")
    major = get_student_leader_data.major

    research_title_list = []

    defense_date = get_panel_invitation.research_title_defense_date
    defense_start_time = get_panel_invitation.research_title_defense_start_time
    defense_end_time = get_panel_invitation.research_title_defense_end_time

    dit_head_full_name = get_panel_invitation.dit_head_full_name
    dit_head_response = get_panel_invitation.dit_head_response
    dit_head_response_date = get_panel_invitation.dit_head_response_date

    panel_full_name = get_panel_invitation.panel_full_name
    panel_username = get_panel_invitation.panel_username
    panel_response = get_panel_invitation.panel_response
    panel_response_date = get_panel_invitation.panel_response_date
    ############## BET-3 PANEL INVITATION DATA ##############

    if get_group_members:
        for group_member in get_group_members:
            student_member_list.append(group_member.student_member_full_name)

    student_member_list.sort()

    if get_research_titles:
        for research_title in get_research_titles:
            research_title_list.append(research_title.research_title)

    print("Date Sent: ", date_submitted)

    print("Group Members: ", student_member_list)
    print("Course: ", course)
    print("Major: ", major)
    print("Research Titles: ", research_title_list)

    print("Defense Date: ", defense_date)
    print("Defense Start Time: ", defense_start_time)
    print("Defense End Time: ", defense_end_time)

    print("DIT Head Name: ", dit_head_full_name)
    print("DIT Head Response: ", dit_head_response)
    print("DIT Head Response Date: ", dit_head_response_date)

    print("Panel Name: ", panel_full_name)
    print("Panel Response: ", panel_response)
    print("Panel Response Date: ", panel_response_date)

    # Local Host - Syntax
    doc = Document("static/forms/1-TOPIC-DEFENSE-PANEL-INVITATION.docx")

    # Deployed - Syntax
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/1-TOPIC-DEFENSE-PANEL-INVITATION.docx')

    student_table = doc.tables[1]
    dit_sign_box = doc.tables[2]
    panel_sign_box = doc.tables[3]
    panel_table = doc.tables[4]
    qr_code_box = doc.tables[5]

    try:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = student_member_list[0]
        student_table.cell(1, 2).paragraphs[0].runs[0].text = course
        student_table.cell(1, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = student_member_list[1]
        student_table.cell(2, 2).paragraphs[0].runs[0].text = course
        student_table.cell(2, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = student_member_list[2]
        student_table.cell(3, 2).paragraphs[0].runs[0].text = course
        student_table.cell(3, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = student_member_list[3]
        student_table.cell(4, 2).paragraphs[0].runs[0].text = course
        student_table.cell(4, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = student_member_list[4]
        student_table.cell(5, 2).paragraphs[0].runs[0].text = course
        student_table.cell(5, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 4).paragraphs[0].runs[0].text = ""

    doc.paragraphs[0].runs[3].text = date_submitted
    doc.paragraphs[1].runs[2].text = panel_full_name

    try:
        doc.paragraphs[8].runs[0].text = research_title_list[0]
    except:
        doc.paragraphs[8].runs[0].text = ""

    try:
        doc.paragraphs[9].runs[0].text = research_title_list[1]
    except:
        doc.paragraphs[9].runs[0].text = ""

    try:
        doc.paragraphs[10].runs[0].text = research_title_list[2]
    except:
        doc.paragraphs[10].runs[0].text = ""

    try:
        doc.paragraphs[11].runs[0].text = research_title_list[3]
    except:
        doc.paragraphs[11].runs[0].text = ""

    try:
        doc.paragraphs[12].runs[0].text = research_title_list[4]
    except:
        doc.paragraphs[12].runs[0].text = ""

    doc.paragraphs[14].runs[1].text = defense_date
    doc.paragraphs[14].runs[4].text = defense_start_time
    doc.paragraphs[14].runs[6].text = defense_end_time
    doc.paragraphs[22].runs[0].text = dit_head_full_name

    panel_table.cell(0, 9).paragraphs[0].runs[0].text = panel_response_date

    if panel_response == "accepted":
        panel_table.cell(0, 1).paragraphs[0].runs[0].text = "✓"
        panel_table.cell(0, 3).paragraphs[0].runs[0].text = "__"

    if panel_response == "declined":
        panel_table.cell(0, 1).paragraphs[0].runs[0].text = "__"
        panel_table.cell(0, 3).paragraphs[0].runs[0].text = "✓"


    if get_panel_invitation.dit_head_signature == 1:
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_invitation.dit_head_username) + ".png"):
            dit_sign_box.cell(0, 0).text = ''
            head_signature = dit_sign_box.cell(0, 0).add_paragraph()
            head_signature_run = head_signature.add_run()
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_invitation.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            get_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username)
            get_accepted_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
            get_pending_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_invitations": get_panel_invitations,
                "accepted_panel_invitations": get_accepted_panel_invitations.count(),
                "pending_panel_invitations": get_pending_panel_invitations.count(),
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-panel-invitation-dashboard.html", context)

    if get_panel_invitation.panel_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_invitation.panel_username) + ".png"):
            panel_sign_box.cell(0, 0).text = ''
            panel_signature = panel_sign_box.cell(0, 0).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_invitation.panel_username +'.png',width=Inches(1.2), height=Inches(0.45))


        else:
            get_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username)
            get_accepted_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
            get_pending_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_invitations": get_panel_invitations,
                "accepted_panel_invitations": get_accepted_panel_invitations.count(),
                "pending_panel_invitations": get_pending_panel_invitations.count(),
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-panel-invitation-dashboard.html", context)


    # Create - QR Code
    auth_qr_code = qrcode.make('Topic Panel Invitation\nH:' + dit_head_response_date + '\nP: ' + panel_response_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-TOPIC-DEFENSE-PANEL-INVITATION-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-TOPIC-DEFENSE-PANEL-INVITATION-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-DEFENSE-PANEL-INVITATION.docx")
    convert(current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-DEFENSE-PANEL-INVITATION.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-DEFENSE-PANEL-INVITATION.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-DEFENSE-PANEL-INVITATION.pdf")

    # Deployed - Syntax
    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-DEFENSE-PANEL-INVITATION.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-TOPIC-DEFENSE-PANEL-INVITATION.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-DEFENSE-PANEL-INVITATION.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-DEFENSE-PANEL-INVITATION.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-TOPIC-DEFENSE-PANEL-INVITATION-QR.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Panel Invitation for Topic Defense QR - Deleted")
    else:
        print("Panel Invitation for Topic Defense QR - Not Found")

    bet3_topic_defense_panel_inviation_docx = current_user.username+"-"+panel_username+"-"+panel_response+"-TOPIC-DEFENSE-PANEL-INVITATION.docx"
    
    # Deployed - Syntax
    # bet3_topic_defense_panel_inviation_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-TOPIC-DEFENSE-PANEL-INVITATION.docx')

    if os.path.isfile(bet3_topic_defense_panel_inviation_docx):
        os.remove(bet3_topic_defense_panel_inviation_docx)
        print("Panel Invitation for Topic Defense docx - Deleted")
    else:
        print("Panel Invitation for Topic Defense docx - Not Found")

    get_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username)
    get_accepted_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
    get_pending_panel_invitations = TitlePanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        #'download_link': download_link,
        "panel_invitations": get_panel_invitations,
        "accepted_panel_invitations": get_accepted_panel_invitations.count(),
        "pending_panel_invitations": get_pending_panel_invitations.count(),
        "response": "sweet downloaded",
    }

    return render(request, "student-bet3-panel-invitation-dashboard.html", context)


# Student - BET3 - Critique Form - Download
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentTopicCritiqueFormDownload(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    # Student - Get Student Leader Data
    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
    except:
        return redirect("login")
    
    # Get Student Group Members
    try:
        get_student_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
    except:
        get_student_group_members = None

    try:
        research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, title_defense_status = "Accepted")
    except:
        pass
    try:
        research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, title_defense_status = "Revise Title")
    except:
        pass

    if get_student_leader_data.middle_name == "":
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name
    else:
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name + " " + get_student_leader_data.middle_name[0] + "."

    get_critique_form = TitleDefenseCritique.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_title_defense_date)
    get_proposal_defense_form = TitleDefenseForm.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_title_defense_date)

    get_critique_form_panel_data = TitleDefenseCritique.objects.all().filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = False, defense_date = get_student_leader_data.research_title_defense_date)
    get_critique_form_panel_chair_data = TitleDefenseCritique.objects.filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = True, defense_date = get_student_leader_data.research_title_defense_date)

    doc = Document("static/forms/5-CRITIQUE-FORM.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/5-CRITIQUE-FORM.docx')
    
    paragraph = doc.add_paragraph()

    for i in range(len(get_critique_form)):
        if get_critique_form[i].critique:
            paragraph.add_run("     ● " + get_critique_form[i].critique)
            run = paragraph.add_run()
            run.add_break()
        i + 1
    

    panel_1 = doc.tables[1]
    panel_chair = doc.tables[2]
    panel_3 = doc.tables[3]
    panel_4 = doc.tables[5]
    panel_5 = doc.tables[6]
    panel_table_1 = doc.tables[4]
    panel_table_2 = doc.tables[8]
    date_table = doc.tables[7]
    qr_code_box = doc.tables[9]

    # Date of the Form
    date_table.cell(0, 1).text = get_critique_form_panel_chair_data[0].defense_date

    # Panel Chairman
    if get_critique_form_panel_chair_data[0]:
        panel_table_1.cell(1, 2).paragraphs[0].runs[0].text = get_critique_form_panel_chair_data[0].panel_full_name
        if get_critique_form_panel_chair_data[0].panel_chairman_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_chair_data[0].panel_username) + ".png"):
                panel_table_1.cell(0, 2).text = ""
                panel_chair_sign = panel_chair.cell(0, 0).add_paragraph()
                panel_chair_sign_run = panel_chair_sign.add_run()
                panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": research_title.research_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-topic-critique-form.html", context)

    # Panel 1
    if get_critique_form_panel_data[0]:
        panel_table_1.cell(1, 0).paragraphs[0].runs[0].text = get_critique_form_panel_data[0].panel_full_name
        if get_critique_form_panel_data[0].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[0].panel_username) + ".png"):
                panel_table_1.cell(0, 0).text = ""
                panel_sign = panel_1.cell(0, 0).add_paragraph()
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": research_title.research_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-topic-critique-form.html", context)
    
    # Panel 2
    if get_critique_form_panel_data[1]:
        panel_table_1.cell(1, 4).paragraphs[0].runs[0].text = get_critique_form_panel_data[1].panel_full_name
        if get_critique_form_panel_data[1].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[1].panel_username) + ".png"):
                panel_table_1.cell(0, 4).text = ""
                panel_sign = panel_3.cell(0, 0).add_paragraph()
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[1].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": research_title.research_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-topic-critique-form.html", context)
    
    # Panel 3
    try:
        if get_critique_form_panel_data[2]:
            panel_table_2.cell(1, 0).paragraphs[0].runs[0].text = get_critique_form_panel_data[2].panel_full_name
            if get_critique_form_panel_data[2].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[2].panel_username) + ".png"):
                    panel_table_2.cell(0, 0).text = ""
                    panel_sign = panel_4.cell(0, 0).add_paragraph()
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[2].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                    # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": research_title.research_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-topic-critique-form.html", context)
    except:
        panel_table_2.cell(1, 0).paragraphs[0].runs[0].text = ""
        panel_table_2.cell(0, 0).text = ""

    try:
        # Panel 4
        if get_critique_form_panel_data[3]:
            panel_table_2.cell(1, 4).paragraphs[0].runs[0].text = get_critique_form_panel_data[3].panel_full_name
            if get_critique_form_panel_data[1].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[3].panel_username) + ".png"):
                    panel_table_2.cell(0, 4).text = ""
                    panel_sign = panel_5.cell(0, 0).add_paragraph()
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[3].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                    # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": research_title.research_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-topic-critique-form.html", context)
    except:
        panel_table_2.cell(1, 4).paragraphs[0].runs[0].text = ""
        panel_table_2.cell(0, 4).text = ""


    # Create - QR Code
    auth_qr_code = qrcode.make('Topic Critique Form\nDate:' + get_student_leader_data.research_title_defense_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-BET3-CRITIQUE-FORM-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET3-CRITIQUE-FORM-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-TOPIC-CRITIQUE-FORM.docx")
    convert(current_user.username + "-TOPIC-CRITIQUE-FORM.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-TOPIC-CRITIQUE-FORM.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-TOPIC-CRITIQUE-FORM.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-CRITIQUE-FORM.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-TOPIC-CRITIQUE-FORM.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-CRITIQUE-FORM.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-CRITIQUE-FORM.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET3-CRITIQUE-FORM-QR.png"

    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Critique Form QR - Deleted")
    else:
        print("Critique Form QR - Not Found")

    delete_critique_form = current_user.username + "-TOPIC-CRITIQUE-FORM.docx"
    # delete_critique_form = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-CRITIQUE-FORM.docx')

    if os.path.isfile(delete_critique_form):
        os.remove(delete_critique_form)
        print("Critique Form - Deleted")
    else:
        print("Critique Form - Form")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        "student_leader_full_name": student_leader_full_name,
        "student_group_members": get_student_group_members,
        "student_research_title": research_title.research_title,
        "critiques": get_critique_form,
        "proposal_defense_form": get_proposal_defense_form,
        "response": "sweet downloaded"
    }

    return render(request, "student-topic-critique-form.html", context)


# Student - BET3 - Research Title Defense - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentDownloadBET3ResearchTitleDefenseForm(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_panel_chairman = TitleDefenseForm.objects.get(student_leader_username=current_user.username, is_panel_chairman=True)
        print("pass 1")
    except:
        print("pass 1.1")
        return redirect("student-bet3-research-title-defense")

    get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
    get_research_titles = ResearchTitle.objects.all().filter(student_leader_username=current_user.username)
    get_panel_members = TitleDefenseForm.objects.all().filter(student_leader_username=current_user.username, is_panel_chairman=False)

    ############## BET-3 RESEARCH TITLE DEFENSE DATA ##############
    date = get_student_leader_data.research_title_defense_date

    try:
        student_1 = get_research_titles[0].student_leader_name
        degree_1 = get_student_leader_data.course_major_abbr
    except:
        student_1 = ""
        degree_1 = ""

    try:
        student_2 = get_group_members[0].student_member_full_name
        degree_2 = get_student_leader_data.course_major_abbr
    except:
        student_2 = ""
        degree_2 = ""

    try:
        student_3 = get_group_members[1].student_member_full_name
        degree_3 = get_student_leader_data.course_major_abbr
    except:
        student_3 = ""
        degree_3 = ""

    try:
        student_4 = get_group_members[2].student_member_full_name
        degree_4 = get_student_leader_data.course_major_abbr
    except:
        student_4 = ""
        degree_4 = ""

    try:
        student_5 = get_group_members[3].student_member_full_name
        degree_5 = get_student_leader_data.course_major_abbr
    except:
        student_5 = ""
        degree_5 = ""

    accepted_title = None
    revise_title = None
    suggested_title = None

    try:
        if get_research_titles[0].old_research_title:
            title_1 = get_research_titles[0].old_research_title
        else:
             title_1 = get_research_titles[0].research_title

        if get_research_titles[0].title_defense_status == "Accepted":
            comment_1 = "Accepted"
            accepted_title = get_research_titles[0].research_title
        elif get_research_titles[0].title_defense_status == "Revise Title":
            comment_1 = "Revise Title"
            revise_title = get_research_titles[0].research_title
        elif get_research_titles[0].title_defense_status == "Deferred":
            comment_1 = "Deferred"

    except:
        title_1 = ""
        comment_1 = ""

    try:
        title_2 = get_research_titles[1].research_title

        if get_research_titles[1].title_defense_status == "Accepted":
            comment_2 = "Accepted"
            accepted_title = get_research_titles[1].research_title
        elif get_research_titles[1].title_defense_status == "Revise Title":
            comment_2 = "Revise Title"
            revise_title = get_research_titles[1].research_title
        elif get_research_titles[1].title_defense_status == "Deferred":
            comment_2 = "Deferred"

    except:
        title_2 = ""
        comment_2 = ""

    try:
        title_3 = get_research_titles[2].research_title

        if get_research_titles[2].title_defense_status == "Accepted":
            comment_3 = "Accepted"
            accepted_title = get_research_titles[2].research_title
        elif get_research_titles[2].title_defense_status == "Revise Title":
            revise_title = get_research_titles[2].research_title
        elif get_research_titles[2].title_defense_status == "Deferred":
            comment_3 = "Deferred"

    except:
        title_3 = ""
        comment_3 = ""

    try:
        title_4 = get_research_titles[3].research_title

        if get_research_titles[3].title_defense_status == "Accepted":
            comment_4 = "Accepted"
            accepted_title = get_research_titles[3].research_title
        elif get_research_titles[3].title_defense_status == "Revise Title":
            comment_4 = "Revise Title"
            revise_title = get_research_titles[3].research_title
        elif get_research_titles[3].title_defense_status == "Deferred":
            comment_4 = "Deferred"

    except:
        title_4 = ""
        comment_4 = ""

    try:
        title_5 = get_research_titles[4].research_title

        if get_research_titles[4].title_defense_status == "Accepted":
            comment_5 = "Accepted"
            accepted_title = get_research_titles[4].research_title
        elif get_research_titles[4].title_defense_status == "Revise Title":
            comment_5 = "Revise Title"
            revise_title = get_research_titles[4].research_title
        elif get_research_titles[4].title_defense_status == "Deferred":
            comment_5 = "Deferred"

    except:
        title_5 = ""
        comment_5 = ""

    try:
        if get_research_titles[0].suggested_title != "":
            suggested_title = get_research_titles[0].suggested_title
    except:
        pass

    try:
        if get_research_titles[1].suggested_title != "":
            suggested_title = get_research_titles[1].suggested_title
    except:
        pass

    try:
        if get_research_titles[2].suggested_title != "":
            suggested_title = get_research_titles[2].suggested_title
    except:
        pass

    try:
        if get_research_titles[3].suggested_title != "":
            suggested_title = get_research_titles[3].suggested_title
    except:
        pass

    try:
        if get_research_titles[4].suggested_title != "":
            suggested_title = get_research_titles[4].suggested_title
    except:
        pass

    if not suggested_title:
        suggested_title = "___________________________________"

    panel_chairman = get_panel_chairman.panel_full_name

    try:
        panel_1 = get_panel_members[0].panel_full_name
    except:
        panel_1 = ""

    try:
        panel_2 = get_panel_members[1].panel_full_name
    except:
        panel_2 = ""

    try:
        panel_3 = get_panel_members[2].panel_full_name
    except:
        panel_3 = ""

    try:
        panel_4 = get_panel_members[3].panel_full_name
    except:
        panel_4 = ""

    panel_5 = ""
    panel_6 = ""
    ############## BET-3 RESEARCH TITLE DEFENSE DATA ##############

    doc = Document("static/forms/2-RESEARCH-TITLE-DEFENSE-FORM.docx")

    # Deployed - Syntax
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/2-RESEARCH-TITLE-DEFENSE-FORM.docx')

    doc.paragraphs[0].runs[1].text = date
    print(doc.paragraphs[0].runs[1].text)  # Date of Title defense

    # EXAMINEE
    student_name = doc.tables[1]
    student_name.cell(1, 3).paragraphs[0].runs[0].text = student_1
    student_name.cell(2, 3).paragraphs[0].runs[0].text = student_2
    student_name.cell(3, 3).paragraphs[0].runs[0].text = student_3
    student_name.cell(4, 3).paragraphs[0].runs[0].text = student_4
    student_name.cell(5, 3).paragraphs[0].runs[0].text = student_5
    print(student_name.cell(1, 3).text)
    print(student_name.cell(2, 3).text)
    print(student_name.cell(3, 3).text)
    print(student_name.cell(4, 3).text)
    print(student_name.cell(5, 3).text)
    # column - row

    # STUDENT COURSE
    student_course = doc.tables[1]
    student_course.cell(1, 5).paragraphs[0].runs[0].text = degree_1
    student_course.cell(2, 5).paragraphs[0].runs[0].text = degree_2
    student_course.cell(3, 5).paragraphs[0].runs[0].text = degree_3
    student_course.cell(4, 5).paragraphs[0].runs[0].text = degree_4
    student_course.cell(5, 5).paragraphs[0].runs[0].text = degree_5
    print(student_course.cell(1, 5).text)
    print(student_course.cell(2, 5).text)
    print(student_course.cell(3, 5).text)
    print(student_course.cell(4, 5).text)
    print(student_course.cell(5, 5).text)

    # TITLES
    titles = doc.tables[2]
    titles.cell(1, 1).paragraphs[0].runs[0].text = title_1
    titles.cell(2, 1).paragraphs[0].runs[0].text = title_2
    titles.cell(3, 1).paragraphs[0].runs[0].text = title_3
    titles.cell(4, 1).paragraphs[0].runs[0].text = title_4
    titles.cell(5, 1).paragraphs[0].runs[0].text = title_5
    print(titles.cell(1, 1).text)
    print(titles.cell(2, 1).text)
    print(titles.cell(3, 1).text)
    print(titles.cell(4, 1).text)
    print(titles.cell(5, 1).text)

    # COMMENT
    comment = doc.tables[2]
    comment.cell(1, 2).paragraphs[0].runs[0].text = comment_1
    comment.cell(2, 2).paragraphs[0].runs[0].text = comment_2
    comment.cell(3, 2).paragraphs[0].runs[0].text = comment_3
    comment.cell(4, 2).paragraphs[0].runs[0].text = comment_4
    comment.cell(5, 2).paragraphs[0].runs[0].text = comment_5
    print(comment.cell(1, 2).text)
    print(comment.cell(2, 2).text)
    print(comment.cell(3, 2).text)
    print(comment.cell(4, 2).text)
    print(comment.cell(5, 2).text)

    doc.paragraphs[2].runs[1].text = suggested_title
    print(doc.paragraphs[2].runs[1].text)  # SUGGESTEDTITLE

   

    doc.paragraphs[5].runs[0].text = panel_chairman
    print(doc.paragraphs[6].runs[0].text)  # PANELCHAIRMAN

    # PANEL SIGNATURE
    panel123 = doc.tables[7]
    panel123.cell(0, 0).paragraphs[0].runs[0].text = panel_1
    panel123.cell(0, 2).paragraphs[0].runs[0].text = panel_2
    panel123.cell(0, 4).paragraphs[0].runs[0].text = panel_3

    panel456 = doc.tables[9]
    panel456.cell(0, 0).paragraphs[0].runs[0].text = panel_4
    # panel456.cell(0, 2).paragraphs[0].runs[0].text = ""
    # panel456.cell(0, 4).paragraphs[0].runs[0].text = ""

    panel_chair_signature_table = doc.tables[3]

    if get_panel_chairman.panel_signature_attach == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_chairman.panel_username) + ".png"):
            panel_chair_signature_table.cell(0, 0).text = ''
            panel_signature = panel_chair_signature_table.cell(0, 0).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_chairman.panel_username +'.png',width=Inches(1.2), height=Inches(0.45))

        else:
            context = {
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-research-title-defense.html", context)
    

    # Panel 1 Signature Box
    panel1_sign_box = doc.tables[5]
    print(panel1_sign_box.cell(0, 0).text)

    try:
        panel1_sign = get_panel_members[0].panel_full_name
        print("pass")
        if get_panel_members[0].panel_signature_attach == True:
            print("pass")
            # Check if Panel E-Sign Exist
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_members[0].panel_username) + ".png"):
                panel1_sign_box.cell(0, 0).text = ''
                panel_signature = panel1_sign_box.cell(0, 0).add_paragraph()
                panel_signature_run = panel_signature.add_run()
                panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_members[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                if get_panel_chairman.panel_signature_attach == False:
                    panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_members[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "response": "sweet faculty member no signature",
                }

                return render(request, "student-bet3-research-title-defense.html", context)
    except:
       panel1_sign_box.cell(0, 0).text = ''

    # Panel 2 Signature Box
    panel2_sign_box = doc.tables[6]
    print(panel1_sign_box.cell(0, 0).text)

    try:
        panel1_sign = get_panel_members[1].panel_full_name
        print("pass")
        if get_panel_members[1].panel_signature_attach == True:
            print("pass")
            # Check if Panel E-Sign Exist
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_members[1].panel_username) + ".png"):
                panel2_sign_box.cell(0, 0).text = ''
                panel_signature = panel2_sign_box.cell(0, 0).add_paragraph()
                panel_signature_run = panel_signature.add_run()
                panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_members[1].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "response": "sweet faculty member no signature",
                }

                return render(request, "student-bet3-research-title-defense.html", context)
    except:
        panel2_sign_box.cell(0, 0).text = ''

    # Panel 3 Signature Box
    panel3_sign_box = doc.tables[4]
    print(panel3_sign_box.cell(0, 0).text)

    try:
        panel1_sign = get_panel_members[2].panel_full_name
        print("pass")
        if get_panel_members[2].panel_signature_attach == True:
            print("pass")
            # Check if Panel E-Sign Exist
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_members[2].panel_username) + ".png"):
                panel3_sign_box.cell(0, 0).text = ''
                panel_signature = panel3_sign_box.cell(0, 0).add_paragraph()
                panel_signature_run = panel_signature.add_run()
                panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_members[2].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "response": "sweet faculty member no signature",
                }

                return render(request, "student-bet3-research-title-defense.html", context)
    except:
        panel3_sign_box.cell(0, 0).text = ''


    # Panel 4 Signature Box
    panel4_sign_box = doc.tables[8]
    print(panel4_sign_box.cell(0, 0).text)

    try:
        panel1_sign = get_panel_members[3].panel_full_name
        print("pass")
        if get_panel_members[3].panel_signature_attach == True:
            print("pass")
            # Check if Panel E-Sign Exist
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_members[3].panel_username) + ".png"):
                panel4_sign_box.cell(0, 0).text = ''
                panel_signature = panel4_sign_box.cell(0, 0).add_paragraph()
                panel_signature_run = panel_signature.add_run()
                panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_members[3].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "response": "sweet faculty member no signature",
                }

                return render(request, "student-bet3-research-title-defense.html", context)
    except:
       panel4_sign_box.cell(0, 0).text = ''

    if accepted_title:
        img = qrcode.make(
            "Research Title Defense"
            + "\n Date & Time of Defense:"
            + get_student_leader_data.research_title_defense_date
            + " "
            + get_student_leader_data.research_title_defense_start_time
            + " to "
            + get_student_leader_data.research_title_defense_end_time
        )
        type(img)
        img.save(current_user.username + "-BET3-TITLE-DEFENSE-QR.png")

    if revise_title:
        img = qrcode.make(
            "Research Title Defense"
            + "\n Date & Time of Defense:"
            + get_student_leader_data.research_title_defense_date
            + " "
            + get_student_leader_data.research_title_defense_start_time
            + " to "
            + get_student_leader_data.research_title_defense_end_time
        )
        type(img)
        img.save(current_user.username + "-BET3-TITLE-DEFENSE-QR.png")

    qr_code_box = doc.tables[10]

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET3-TITLE-DEFENSE-QR.png", width=Inches(1), height=Inches(1))
    # INSERT IMAGE

    # SAVE DOCX - COVERT TO PDF
    doc.save(current_user.username + "-BET3-RESEARCH-TITLE-DEFENSE-FORM.docx")
    convert(current_user.username + "-BET3-RESEARCH-TITLE-DEFENSE-FORM.docx")
    # SAVE DOCX - COVERT TO PDF

    filePath = FilePath(student_leader_username=current_user.username, file_path=current_user.username + "-BET3-RESEARCH-TITLE-DEFENSE-FORM.pdf")
    filePath.save()

    os.startfile(current_user.username + "-BET3-RESEARCH-TITLE-DEFENSE-FORM.pdf")

    # Deployed - Syntax
    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username + '-BET3-RESEARCH-TITLE-DEFENSE-FORM.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username + '-BET3-RESEARCH-TITLE-DEFENSE-FORM.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/"+current_user.username + '-BET3-RESEARCH-TITLE-DEFENSE-FORM.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username + '-BET3-RESEARCH-TITLE-DEFENSE-FORM.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET3-TITLE-DEFENSE-QR.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Research Title Defense QR Code - Deleted")
    else:
        print("Research Title Defense QR Code - Not Found")

    bet3_research_title_defense_form_docx = current_user.username + "-BET3-RESEARCH-TITLE-DEFENSE-FORM.docx"

    # Deployed - Syntax
    # bet3_research_title_defense_form_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username + '-BET3-RESEARCH-TITLE-DEFENSE-FORM.docx')

    if os.path.isfile(bet3_research_title_defense_form_docx):
        os.remove(bet3_research_title_defense_form_docx)
        print("Research Title Defense Form - Deleted")
    else:
        print("Research Title Defense Form - Not Found")

    # Get Student Group Members
    try:
        get_student_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
    except:
        get_student_group_members = None

    # Get Student Research Title / Titles
    try:
        get_student_research_titles = ResearchTitle.objects.all().filter(student_leader_username=current_user.username)
    except:
        print("pass research titles")
        return redirect("student-dashboard")

    # Get Student Accepted Research Title
    try:
        get_student_accepted_research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, status="Title Defense - Accepted")
    except:
        get_student_accepted_research_title = None

    # Get Student Accepted - Revise Title Research Title
    try:
        get_student_revise_research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, status="Title Defense - Revise Title")
    except:
        get_student_revise_research_title = None

    # Get Panel Research Title Defense Form
    try:
        get_panel_research_title_defense_form = TitleDefenseForm.objects.all().filter(student_leader_username=current_user.username)
    except:
        print("pass research title defense form")
        return redirect("student-dashboard")

    if get_student_leader_data.group_members_status != "completed" and get_student_leader_data.research_titles_status != "completed" and get_student_leader_data.bet3_panel_invitation_status != "completed":
        return redirect("student-dashboard")

    if get_student_leader_data.middle_name == "":
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name
    else:
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name + " " + get_student_leader_data.middle_name[0] + "."

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        "student_leader_full_name": student_leader_full_name,
        "student_group_members": get_student_group_members,
        "student_research_titles": get_student_research_titles,
        "student_accepted_research_title": get_student_accepted_research_title,
        "student_revise_research_title": get_student_revise_research_title,
        "panel_research_title_defense_form": get_panel_research_title_defense_form,
        #'download_link': download_link,
        "response": "sweet downloaded",
    }

    return render(request, "student-bet3-research-title-defense.html", context)


# Student - BET3 - Topic Defense - Panel Invitation - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentTopicPanelConformeDownload(request, id):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
        get_panel_conforme = PanelConforme.objects.get(id=int(id), student_leader_username=current_user.username)

    except:
        return redirect("student-bet3-topic-panel-conforme")

    ############## BET-3 PANEL INVITATION DATA ##############
    date_submitted = get_panel_conforme.form_date_sent

    student_member_list = [get_panel_conforme.student_leader_full_name]
    course = get_student_leader_data.course.replace("Engineering", "Eng.")
    major = get_student_leader_data.major

    defense_date = get_panel_conforme.defense_date
    defense_start_time = get_panel_conforme.defense_start_time
    defense_end_time = get_panel_conforme.defense_end_time

    dit_head_full_name = get_panel_conforme.dit_head_full_name
    dit_head_response = get_panel_conforme.dit_head_response
    dit_head_response_date = get_panel_conforme.dit_head_response_date

    panel_full_name = get_panel_conforme.panel_full_name
    panel_username = get_panel_conforme.panel_username
    panel_response = get_panel_conforme.panel_response
    panel_response_date = get_panel_conforme.panel_response_date

    if get_group_members:
        for group_member in get_group_members:
            student_member_list.append(group_member.student_member_full_name)

    student_member_list.sort()

    print("Date Sent: ", date_submitted)

    print("Group Members: ", student_member_list)
    print("Course: ", course)
    print("Major: ", major)
    print("Research Titles: ", get_panel_conforme.research_title)

    print("Defense Date: ", defense_date)
    print("Defense Start Time: ", defense_start_time)
    print("Defense End Time: ", defense_end_time)

    print("DIT Head Name: ", dit_head_full_name)
    print("DIT Head Response: ", dit_head_response)
    print("DIT Head Response Date: ", dit_head_response_date)

    print("Panel Name: ", panel_full_name)
    print("Panel Response: ", panel_response)
    print("Panel Response Date: ", panel_response_date)

    doc = Document("static/forms/TOPIC-PANEL-CONFORME.docx")

    # Deployed - Syntax
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/TOPIC-PANEL-CONFORME.docx')

    student_table = doc.tables[1]

    dit_head_signature_box = doc.tables[2]
    panel_signature_box = doc.tables[4]

    panel_data = doc.tables[5]

    qr_code_box = doc.tables[3]

    try:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = student_member_list[0]
        student_table.cell(1, 2).paragraphs[0].runs[0].text = course
        student_table.cell(1, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = student_member_list[1]
        student_table.cell(2, 2).paragraphs[0].runs[0].text = course
        student_table.cell(2, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = student_member_list[2]
        student_table.cell(3, 2).paragraphs[0].runs[0].text = course
        student_table.cell(3, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = student_member_list[3]
        student_table.cell(4, 2).paragraphs[0].runs[0].text = course
        student_table.cell(4, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = student_member_list[4]
        student_table.cell(5, 2).paragraphs[0].runs[0].text = course
        student_table.cell(5, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 4).paragraphs[0].runs[0].text = ""

    doc.paragraphs[6].runs[5].text = get_student_leader_data.course + " Major in " + get_student_leader_data.major

    doc.paragraphs[0].runs[1].text = date_submitted
    doc.paragraphs[1].runs[0].text = panel_full_name
    doc.paragraphs[4].runs[1].text = panel_full_name

    doc.paragraphs[8].runs[2].text = get_panel_conforme.research_title
    doc.paragraphs[19].runs[0].text = get_panel_conforme.dit_head_full_name

    panel_data.cell(0, 0).text = panel_full_name
    panel_data.cell(0, 2).text = panel_response_date


    if get_panel_conforme.dit_head_signature == 1:
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_conforme.dit_head_username) + ".png"):
            dit_head_signature_box.cell(0, 0).text = ''
            head_signature = dit_head_signature_box.cell(0, 0).add_paragraph()
            head_signature_run = head_signature.add_run()
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_conforme.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            get_all_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Topic Panel Conforme")
            get_all_pending_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Topic Panel Conforme", is_completed = False)

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_conforme": get_all_topic_panel_conforme,
                "pending_panel_conforme": get_all_pending_topic_panel_conforme,
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-panel-conforme-dashboard.html", context)

    if get_panel_conforme.panel_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_conforme.panel_username) + ".png"):
            panel_signature_box.cell(0, 0).text = ''
            panel_signature = panel_signature_box.cell(0, 0).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_conforme.panel_username +'.png',width=Inches(1.2), height=Inches(0.45))


        else:
            get_all_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Topic Panel Conforme")
            get_all_pending_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Topic Panel Conforme", is_completed = False)

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_conforme": get_all_topic_panel_conforme,
                "pending_panel_conforme": get_all_pending_topic_panel_conforme,
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-panel-conforme-dashboard.html", context)


    # Create - QR Code
    auth_qr_code = qrcode.make('Topic Panel Conforme\nH:' + dit_head_response_date + '\nP: ' + panel_response_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-TOPIC-PANEL-CONFORME-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-TOPIC-PANEL-CONFORME-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-PANEL-CONFORME.docx")
    convert(current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-PANEL-CONFORME.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-PANEL-CONFORME.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-PANEL-CONFORME.pdf")

    # Deployed - Syntax
    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-PANEL-CONFORME.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-TOPIC-PANEL-CONFORME.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-PANEL-CONFORME.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-PANEL-CONFORME.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-TOPIC-PANEL-CONFORME-QR.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Panel Conforme for Topic Defense QR - Deleted")
    else:
        print("Panel Conforme for Topic Defense QR - Not Found")

    bet3_topic_defense_panel_inviation_docx = current_user.username + "-" + panel_username + "-" + panel_response + "-TOPIC-PANEL-CONFORME.docx"

    # Deployed - Syntax
    # bet3_topic_defense_panel_inviation_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-PANEL-CONFORME.docx')

    if os.path.isfile(bet3_topic_defense_panel_inviation_docx):
        os.remove(bet3_topic_defense_panel_inviation_docx)
        print("Panel Conforme for Topic Defense - Deleted")
    else:
        print("Panel Conforme for Topic Defense - Not Found")

    get_all_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Topic Panel Conforme")
    get_all_pending_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Topic Panel Conforme", is_completed = False)

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        #'download_link': download_link,
        "topic_panel_conforme": get_all_topic_panel_conforme,
        "pending_panel_conforme": get_all_pending_topic_panel_conforme,
        "response": "sweet downloaded",
    }

    return render(request, "student-bet3-topic-panel-conforme-dashboard.html", context)


# Student - BET3 - Proposal Defense - Panel Invitation - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentProposalPanelConformeDownload(request, id):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
        get_panel_conforme = PanelConforme.objects.get(id=int(id), student_leader_username=current_user.username)

    except:
        return redirect("student-bet3-proposal-panel-conforme")

    ############## BET-3 PANEL INVITATION DATA ##############
    date_submitted = get_panel_conforme.form_date_sent

    student_member_list = [get_panel_conforme.student_leader_full_name]
    course = get_student_leader_data.course.replace("Engineering", "Eng.")
    major = get_student_leader_data.major

    defense_date = get_panel_conforme.defense_date
    defense_start_time = get_panel_conforme.defense_start_time
    defense_end_time = get_panel_conforme.defense_end_time

    dit_head_full_name = get_panel_conforme.dit_head_full_name
    dit_head_response = get_panel_conforme.dit_head_response
    dit_head_response_date = get_panel_conforme.dit_head_response_date

    panel_full_name = get_panel_conforme.panel_full_name
    panel_username = get_panel_conforme.panel_username
    panel_response = get_panel_conforme.panel_response
    panel_response_date = get_panel_conforme.panel_response_date

    if get_group_members:
        for group_member in get_group_members:
            student_member_list.append(group_member.student_member_full_name)

    student_member_list.sort()

    print("Date Sent: ", date_submitted)

    print("Group Members: ", student_member_list)
    print("Course: ", course)
    print("Major: ", major)
    print("Research Titles: ", get_panel_conforme.research_title)

    print("Defense Date: ", defense_date)
    print("Defense Start Time: ", defense_start_time)
    print("Defense End Time: ", defense_end_time)

    print("DIT Head Name: ", dit_head_full_name)
    print("DIT Head Response: ", dit_head_response)
    print("DIT Head Response Date: ", dit_head_response_date)

    print("Panel Name: ", panel_full_name)
    print("Panel Response: ", panel_response)
    print("Panel Response Date: ", panel_response_date)

    doc = Document("static/forms/PROPOSAL-PANEL-CONFORME.docx")

    # Deployed - Syntax
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/PROPOSAL-PANEL-CONFORME.docx.docx')

    student_table = doc.tables[1]

    dit_head_signature_box = doc.tables[2]
    panel_signature_box = doc.tables[4]

    panel_data = doc.tables[5]

    qr_code_box = doc.tables[3]

    try:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = student_member_list[0]
        student_table.cell(1, 2).paragraphs[0].runs[0].text = course
        student_table.cell(1, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = student_member_list[1]
        student_table.cell(2, 2).paragraphs[0].runs[0].text = course
        student_table.cell(2, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = student_member_list[2]
        student_table.cell(3, 2).paragraphs[0].runs[0].text = course
        student_table.cell(3, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = student_member_list[3]
        student_table.cell(4, 2).paragraphs[0].runs[0].text = course
        student_table.cell(4, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = student_member_list[4]
        student_table.cell(5, 2).paragraphs[0].runs[0].text = course
        student_table.cell(5, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 4).paragraphs[0].runs[0].text = ""

    doc.paragraphs[6].runs[5].text = get_student_leader_data.course + " Major in " + get_student_leader_data.major

    doc.paragraphs[0].runs[1].text = date_submitted
    doc.paragraphs[1].runs[0].text = panel_full_name
    doc.paragraphs[4].runs[1].text = panel_full_name

    doc.paragraphs[8].runs[2].text = get_panel_conforme.research_title
    doc.paragraphs[19].runs[0].text = get_panel_conforme.dit_head_full_name

    panel_data.cell(0, 0).text = panel_full_name
    panel_data.cell(0, 2).text = panel_response_date


    if get_panel_conforme.dit_head_signature == 1:
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_conforme.dit_head_username) + ".png"):
            dit_head_signature_box.cell(0, 0).text = ''
            head_signature = dit_head_signature_box.cell(0, 0).add_paragraph()
            head_signature_run = head_signature.add_run()
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_conforme.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            get_all_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Proposal Panel Conforme")
            get_all_pending_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Proposal Panel Conforme", is_completed = False)

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_conforme": get_all_proposal_panel_conforme,
                "pending_panel_conforme": get_all_pending_proposal_panel_conforme,
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-proposal-panel-conforme-dashboard.html", context)

    if get_panel_conforme.panel_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_conforme.panel_username) + ".png"):
            panel_signature_box.cell(0, 0).text = ''
            panel_signature = panel_signature_box.cell(0, 0).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_conforme.panel_username +'.png',width=Inches(1.2), height=Inches(0.45))


        else:
            get_all_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Proposal Panel Conforme")
            get_all_pending_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Proposal Panel Conforme", is_completed = False)

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_conforme": get_all_topic_panel_conforme,
                "pending_panel_conforme": get_all_pending_topic_panel_conforme,
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-panel-conforme-dashboard.html", context)


    # Create - QR Code
    auth_qr_code = qrcode.make('Proposal Panel Conforme\nH:' + dit_head_response_date + '\nP: ' + panel_response_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-PROPOSAL-PANEL-CONFORME-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-PROPOSAL-PANEL-CONFORME-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-" + panel_username + "-" + panel_response + "-PROPOSAL-PANEL-CONFORME.docx")
    convert(current_user.username + "-" + panel_username + "-" + panel_response + "-PROPOSAL-PANEL-CONFORME.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-" + panel_username + "-" + panel_response + "-PROPOSAL-PANEL-CONFORME.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-" + panel_username + "-" + panel_response + "-PROPOSAL-PANEL-CONFORME.pdf")

    # Deployed - Syntax
    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-PROPOSAL-PANEL-CONFORME.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-PROPOSAL-PANEL-CONFORME.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-PROPOSAL-PANEL-CONFORME.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-PROPOSAL-PANEL-CONFORME.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-PROPOSAL-PANEL-CONFORME-QR.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Panel Conforme for Proposal Defense QR - Deleted")
    else:
        print("Panel Conforme for Proposal Defense QR - Not Found")

    bet3_topic_defense_panel_inviation_docx = current_user.username + "-" + panel_username + "-" + panel_response + "-PROPOSAL-PANEL-CONFORME.docx"
    # bet3_topic_defense_panel_inviation_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-PROPOSAL-PANEL-CONFORME.docx')

    if os.path.isfile(bet3_topic_defense_panel_inviation_docx):
        os.remove(bet3_topic_defense_panel_inviation_docx)
        print("Panel Conforme for Proposal Defense - Deleted")
    else:
        print("Panel Conforme for Proposal Defense - Not Found")

    get_all_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Proposal Panel Conforme")
    get_all_pending_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Proposal Panel Conforme", is_completed = False)

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        #'download_link': download_link,
        "proposal_panel_conforme": get_all_proposal_panel_conforme,
        "pending_panel_conforme": get_all_pending_proposal_panel_conforme,
        "response": "sweet downloaded",
    }

    return render(request, "student-bet3-proposal-panel-conforme-dashboard.html", context)



# Student - BET5 - Final Defense - Panel Invitation - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentFinalPanelConformeDownload(request, id):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
        get_panel_conforme = PanelConforme.objects.get(id=int(id), student_leader_username=current_user.username)

    except:
        return redirect("student-bet3-proposal-panel-conforme")

    ############## BET-3 PANEL INVITATION DATA ##############
    date_submitted = get_panel_conforme.form_date_sent

    student_member_list = [get_panel_conforme.student_leader_full_name]
    course = get_student_leader_data.course.replace("Engineering", "Eng.")
    major = get_student_leader_data.major

    defense_date = get_panel_conforme.defense_date
    defense_start_time = get_panel_conforme.defense_start_time
    defense_end_time = get_panel_conforme.defense_end_time

    dit_head_full_name = get_panel_conforme.dit_head_full_name
    dit_head_response = get_panel_conforme.dit_head_response
    dit_head_response_date = get_panel_conforme.dit_head_response_date

    panel_full_name = get_panel_conforme.panel_full_name
    panel_username = get_panel_conforme.panel_username
    panel_response = get_panel_conforme.panel_response
    panel_response_date = get_panel_conforme.panel_response_date

    if get_group_members:
        for group_member in get_group_members:
            student_member_list.append(group_member.student_member_full_name)

    student_member_list.sort()

    print("Date Sent: ", date_submitted)

    print("Group Members: ", student_member_list)
    print("Course: ", course)
    print("Major: ", major)
    print("Research Titles: ", get_panel_conforme.research_title)

    print("Defense Date: ", defense_date)
    print("Defense Start Time: ", defense_start_time)
    print("Defense End Time: ", defense_end_time)

    print("DIT Head Name: ", dit_head_full_name)
    print("DIT Head Response: ", dit_head_response)
    print("DIT Head Response Date: ", dit_head_response_date)

    print("Panel Name: ", panel_full_name)
    print("Panel Response: ", panel_response)
    print("Panel Response Date: ", panel_response_date)

    doc = Document("static/forms/FINAL-PANEL-CONFORME.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/FINAL-PANEL-CONFORME.docx')

    student_table = doc.tables[1]

    dit_head_signature_box = doc.tables[2]
    panel_signature_box = doc.tables[4]

    panel_data = doc.tables[5]

    qr_code_box = doc.tables[3]

    try:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = student_member_list[0]
        student_table.cell(1, 2).paragraphs[0].runs[0].text = course
        student_table.cell(1, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = student_member_list[1]
        student_table.cell(2, 2).paragraphs[0].runs[0].text = course
        student_table.cell(2, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = student_member_list[2]
        student_table.cell(3, 2).paragraphs[0].runs[0].text = course
        student_table.cell(3, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = student_member_list[3]
        student_table.cell(4, 2).paragraphs[0].runs[0].text = course
        student_table.cell(4, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = student_member_list[4]
        student_table.cell(5, 2).paragraphs[0].runs[0].text = course
        student_table.cell(5, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 4).paragraphs[0].runs[0].text = ""

    doc.paragraphs[6].runs[8].text = get_student_leader_data.course + " Major in " + get_student_leader_data.major

    doc.paragraphs[0].runs[1].text = date_submitted
    doc.paragraphs[1].runs[0].text = panel_full_name
    doc.paragraphs[4].runs[1].text = panel_full_name

    doc.paragraphs[8].runs[2].text = get_panel_conforme.research_title
    doc.paragraphs[19].runs[0].text = get_panel_conforme.dit_head_full_name

    panel_data.cell(0, 0).text = panel_full_name
    panel_data.cell(0, 2).text = panel_response_date


    if get_panel_conforme.dit_head_signature == 1:
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_conforme.dit_head_username) + ".png"):
            dit_head_signature_box.cell(0, 0).text = ''
            head_signature = dit_head_signature_box.cell(0, 0).add_paragraph()
            head_signature_run = head_signature.add_run()
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_conforme.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            get_all_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Final Panel Conforme")
            get_all_pending_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Final Panel Conforme", is_completed = False)

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_conforme": get_all_proposal_panel_conforme,
                "pending_panel_conforme": get_all_pending_proposal_panel_conforme,
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet5-final-panel-conforme-dashboard.html", context)

    if get_panel_conforme.panel_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_conforme.panel_username) + ".png"):
            panel_signature_box.cell(0, 0).text = ''
            panel_signature = panel_signature_box.cell(0, 0).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_conforme.panel_username +'.png',width=Inches(1.2), height=Inches(0.45))


        else:
            get_all_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Final Panel Conforme")
            get_all_pending_topic_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Final Panel Conforme", is_completed = False)

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_conforme": get_all_topic_panel_conforme,
                "pending_panel_conforme": get_all_pending_topic_panel_conforme,
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet5-final-panel-conforme-dashboard.html", context)


    # Create - QR Code
    auth_qr_code = qrcode.make('Final Panel Conforme\nH:' + dit_head_response_date + '\nP: ' + panel_response_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-FINAL-PANEL-CONFORME-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-FINAL-PANEL-CONFORME-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-" + panel_username + "-" + panel_response + "-FINAL-PANEL-CONFORME.docx")
    convert(current_user.username + "-" + panel_username + "-" + panel_response + "-FINAL-PANEL-CONFORME.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-" + panel_username + "-" + panel_response + "-FINAL-PANEL-CONFORME.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-" + panel_username + "-" + panel_response + "-FINAL-PANEL-CONFORME.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-FINAL-PANEL-CONFORME.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-FINAL-PANEL-CONFORME.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-FINAL-PANEL-CONFORME.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-FINAL-PANEL-CONFORME.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-FINAL-PANEL-CONFORME-QR.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Panel Conforme for Final Defense QR - Deleted")
    else:
        print("Panel Conforme for Final Defense QR - Not Found")

    bet3_topic_defense_panel_inviation_docx = current_user.username + "-" + panel_username + "-" + panel_response + "-FINAL-PANEL-CONFORME.docx"
    # bet3_topic_defense_panel_inviation_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-TOPIC-PANEL-CONFORME.docx')

    if os.path.isfile(bet3_topic_defense_panel_inviation_docx):
        os.remove(bet3_topic_defense_panel_inviation_docx)
        print("Panel Conform for Final Defense - Deleted")
    else:
        print("Panel Conforme for Final Defense - Not Found")

    get_all_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Final Panel Conforme")
    get_all_pending_proposal_panel_conforme = PanelConforme.objects.all().filter(student_leader_username=current_user.username, form = "Final Panel Conforme", is_completed = False)

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        #'download_link': download_link,
        "proposal_panel_conforme": get_all_proposal_panel_conforme,
        "pending_panel_conforme": get_all_pending_proposal_panel_conforme,
        "response": "sweet downloaded",
    }

    return render(request, "student-bet5-final-panel-conforme-dashboard.html", context)


# Student - BET-3 Adviser Conforme - Download
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentBET3AdviserConformeDownload(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    # Student - Get Student Leader Data
    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
    except:
        return redirect("logout_user")

    # BET-3 - Get Adviser Conforme
    try:
        get_adviser_conforme = AdviserConforme.objects.get(student_leader_username=current_user.username, form_status="Accepted")
    except:
        return redirect("student-bet3-adviser-dashboard")

    get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)

    if get_student_leader_data.middle_name == "":
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name
    else:
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name + " " + get_student_leader_data.middle_name[0] + "."

    try:
        student_1 = student_leader_full_name
        course_1 = get_student_leader_data.course.replace("Engineering", "Eng.")
        major_1 = get_student_leader_data.major
    except:
        student_1 = ""
        course_1 = ""
        major_1 = ""

    try:
        student_2 = get_group_members[0].student_member_full_name
        course_2 = get_student_leader_data.course.replace("Engineering", "Eng.")
        major_2 = get_student_leader_data.major
    except:
        student_2 = ""
        course_2 = ""
        major_2 = ""

    try:
        student_3 = get_group_members[1].student_member_full_name
        course_3 = get_student_leader_data.course.replace("Engineering", "Eng.")
        major_3 = get_student_leader_data.major
    except:
        student_3 = ""
        course_3 = ""
        major_3 = ""

    try:
        student_4 = get_group_members[2].student_member_full_name
        course_4 = get_student_leader_data.course.replace("Engineering", "Eng.")
        major_4 = get_student_leader_data.major
    except:
        student_4 = ""
        course_4 = ""
        major_4 = ""

    try:
        student_5 = get_group_members[3].student_member_full_name
        course_5 = get_student_leader_data.course.replace("Engineering", "Eng.")
        major_5 = get_student_leader_data.major
    except:
        student_5 = ""
        course_5 = ""
        major_5 = ""

    course_major = get_student_leader_data.course + " Major in " + get_student_leader_data.major
    research_title = get_adviser_conforme.research_title
    date_submitted = get_adviser_conforme.form_date_submitted

    dit_head_name = get_adviser_conforme.dit_head_name

    adviser_name = get_adviser_conforme.adviser_name
    adviser_response_date = get_adviser_conforme.adviser_response_date

    doc = Document("static/forms/3-RESEARCH-ADVISER-CONFORME.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/3-RESEARCH-ADVISER-CONFORME.docx')

    doc.paragraphs[0].runs[3].text = date_submitted

    doc.paragraphs[1].runs[2].text = adviser_name

    # doc.paragraphs[6].runs[1].text = adviser_name

    # STUDENTS NAME
    student_table = doc.tables[0]

    student_table.cell(1, 0).paragraphs[0].runs[0].text = student_1
    student_table.cell(2, 0).paragraphs[0].runs[0].text = student_2
    student_table.cell(3, 0).paragraphs[0].runs[0].text = student_3
    student_table.cell(4, 0).paragraphs[0].runs[0].text = student_4
    student_table.cell(5, 0).paragraphs[0].runs[0].text = student_5

    # student_table.cell(1, 2).paragraphs[0].runs[0].text = course
    # student_table.cell(1, 4).paragraphs[0].runs[0].text = major

    # STUDENT COURSE
    student_table.cell(1, 2).paragraphs[0].runs[0].text = course_1
    student_table.cell(2, 2).paragraphs[0].runs[0].text = course_2
    student_table.cell(3, 2).paragraphs[0].runs[0].text = course_3
    student_table.cell(4, 2).paragraphs[0].runs[0].text = course_4
    student_table.cell(5, 2).paragraphs[0].runs[0].text = course_5

    # STUDENT MAJOR

    student_table.cell(1, 4).paragraphs[0].runs[0].text = major_1
    student_table.cell(2, 4).paragraphs[0].runs[0].text = major_2
    student_table.cell(3, 4).paragraphs[0].runs[0].text = major_3
    student_table.cell(4, 4).paragraphs[0].runs[0].text = major_4
    student_table.cell(5, 4).paragraphs[0].runs[0].text = major_5

    doc.paragraphs[8].runs[1].text = research_title

    doc.paragraphs[22].runs[0].text = dit_head_name

    doc.paragraphs[30].runs[0].text = adviser_name
    doc.paragraphs[33].runs[0].text = adviser_response_date

    head_signature_table = doc.tables[1]
    adviser_signature_table = doc.tables[3]

    if get_adviser_conforme.dit_head_signature == 1:
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_adviser_conforme.dit_head_username) + ".png"):
            head_signature_table.cell(0, 0).text = ''
            head_signature = head_signature_table.cell(0, 0).add_paragraph()
            head_signature_run = head_signature.add_run()
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_adviser_conforme.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_adviser_conforme.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            get_advisers = User.objects.all().filter(is_adviser=1)

            # Student = Check Adviser Conforme
            try:
                check_adviser_conforme = AdviserConforme.objects.get(student_leader_username=current_user.username)
            except:
                check_adviser_conforme = None

                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "advisers": get_advisers,
                    "adviser_conforme_data": check_adviser_conforme,
                    #'download_link': download_link,
                    "response": "sweet faculty member no signature",
                }

                return render(request, "student-bet3-adviser-dashboard.html", context)

    if get_adviser_conforme.adviser_signature == 1:
        print("Adviser attach esign")
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_adviser_conforme.adviser_username) + ".png"):
            print("Adviser esign found")
            adviser_signature_table.cell(0, 0).text = ''
            adviser_signature = adviser_signature_table.cell(0, 0).add_paragraph()
            adviser_signature_run = adviser_signature.add_run()
            adviser_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_adviser_conforme.adviser_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            print("Not Found")
            get_advisers = User.objects.all().filter(is_adviser=1)

            # Student = Check Adviser Conforme
            try:
                check_adviser_conforme = AdviserConforme.objects.get(student_leader_username=current_user.username)
            except:
                check_adviser_conforme = None

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "advisers": get_advisers,
                "adviser_conforme_data": check_adviser_conforme,
                #'download_link': download_link,
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-adviser-dashboard.html", context)

    img = qrcode.make("Research Adviser Conforme: " + "\n H: " + get_adviser_conforme.dit_head_response_date + "\n A: " + get_adviser_conforme.adviser_response_date)
    type(img)
    img.save(current_user.username + "-BET3-ADVISER-CONFORME.png")

    qr_code_box = doc.tables[2]

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET3-ADVISER-CONFORME.png", width=Inches(1), height=Inches(1))
    # INSERT IMAGE

    # SAVE DOCX - COVERT TO PDF
    doc.save(current_user.username + "-BET3-ADVISER-CONFORME.docx")
    convert(current_user.username + "-BET3-ADVISER-CONFORME.docx")
    # SAVE DOCX - COVERT TO PDF

    os.startfile(current_user.username + "-BET3-ADVISER-CONFORME.pdf")

    filePath = FilePath(student_leader_username=current_user.username, file_path=current_user.username + "-BET3-ADVISER-CONFORME.pdf")
    filePath.save()

    # UN COMMENT IF DEPLOYED
    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +'-BET3-ADVISER-CONFORME.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username +'-BET3-ADVISER-CONFORME.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/"+current_user.username +'-BET3-ADVISER-CONFORME.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +'-BET3-ADVISER-CONFORME.pdf'
    # )
    # filePath.save()
    # UN COMMENT IF DEPLOYED

    qr_code_path = current_user.username + "-BET3-ADVISER-CONFORME.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Adviser Conforme QR - Deleted")
    else:
        print("Adviser Conforme QR - Not Found")

    bet3_adviser_conforme_docx = current_user.username + "-BET3-ADVISER-CONFORME.docx"
    # bet3_adviser_conforme_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +'-BET3-ADVISER-CONFORME.pdf')

    if os.path.isfile(bet3_adviser_conforme_docx):
        os.remove(bet3_adviser_conforme_docx)
        print("Adviser Conforme - Deleted")
    else:
        print("Adviser Conforme - Not Found")

    # Student = Check Adviser Conforme
    try:
        check_adviser_conforme = AdviserConforme.objects.get(student_leader_username=current_user.username)
    except:
        check_adviser_conforme = None

    get_advisers = User.objects.all().filter(is_adviser=1)

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        "advisers": get_advisers,
        "adviser_conforme_data": check_adviser_conforme,
        #'download_link': download_link,
        "response": "sweet downloaded",
    }

    return render(request, "student-bet3-adviser-dashboard.html", context)


# Student - BET3 - Proposal Defense - Panel Invitation - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentBET3ProposalDefensePanelInvitationDownload(request, id):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
        get_panel_invitation = ProposalPanelInvitation.objects.get(id=int(id), student_leader_username=current_user.username)

    except:
        return redirect("student-bet3-proposal-defense-panel-invitation-dashboard")
    
    try:
        get_accepted_research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, title_defense_status = "Accepted")
        research_title = get_accepted_research_title.research_title
    except:
        pass

    try:
        get_revise_research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, title_defense_status = "Revise Title")
        research_title = get_revise_research_title.research_title
    except:
        pass


    ############## BET-3 PANEL INVITATION DATA ##############
    date_submitted = get_panel_invitation.form_date_sent

    student_member_list = [get_panel_invitation.student_leader_full_name]
    course = get_student_leader_data.course.replace("Engineering", "Eng.")
    major = get_student_leader_data.major


    defense_date = get_panel_invitation.research_proposal_defense_date
    defense_start_time = get_panel_invitation.research_proposal_defense_start_time
    defense_end_time = get_panel_invitation.research_proposal_defense_end_time

    dit_head_full_name = get_panel_invitation.dit_head_full_name
    dit_head_response = get_panel_invitation.dit_head_response
    dit_head_response_date = get_panel_invitation.dit_head_response_date

    panel_full_name = get_panel_invitation.panel_full_name
    panel_username = get_panel_invitation.panel_username
    panel_response = get_panel_invitation.panel_response
    panel_response_date = get_panel_invitation.panel_response_date
    ############## BET-3 PANEL INVITATION DATA ##############

    if get_group_members:
        for group_member in get_group_members:
            student_member_list.append(group_member.student_member_full_name)

    student_member_list.sort()

    print("Date Sent: ", date_submitted)

    print("Group Members: ", student_member_list)
    print("Course: ", course)
    print("Major: ", major)
    print("Research Title: ", research_title)

    print("Defense Date: ", defense_date)
    print("Defense Start Time: ", defense_start_time)
    print("Defense End Time: ", defense_end_time)

    print("DIT Head Name: ", dit_head_full_name)
    print("DIT Head Response: ", dit_head_response)
    print("DIT Head Response Date: ", dit_head_response_date)

    print("Panel Name: ", panel_full_name)
    print("Panel Response: ", panel_response)
    print("Panel Response Date: ", panel_response_date)

    doc = Document("static/forms/4-PROPOSAL-DEFENSE-PANEL-INVITATION.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/4-PROPOSAL-DEFENSE-PANEL-INVITATION.docx')

    student_table = doc.tables[1] # Student Data Table
    head_signature_table = doc.tables[2]
    panel_table = doc.tables[4] # Panel Data Table
    qr_code_box = doc.tables[5]
    panel_signature_box = doc.tables[3]

    try:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = student_member_list[0]
        student_table.cell(1, 2).paragraphs[0].runs[0].text = course
        student_table.cell(1, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = student_member_list[1]
        student_table.cell(2, 2).paragraphs[0].runs[0].text = course
        student_table.cell(2, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = student_member_list[2]
        student_table.cell(3, 2).paragraphs[0].runs[0].text = course
        student_table.cell(3, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = student_member_list[3]
        student_table.cell(4, 2).paragraphs[0].runs[0].text = course
        student_table.cell(4, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = student_member_list[4]
        student_table.cell(5, 2).paragraphs[0].runs[0].text = course
        student_table.cell(5, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 4).paragraphs[0].runs[0].text = ""

    doc.paragraphs[0].runs[1].text = date_submitted
    doc.paragraphs[1].runs[1].text = panel_full_name

    doc.paragraphs[4].runs[6].text = research_title

    doc.paragraphs[6].runs[1].text = defense_date
    doc.paragraphs[6].runs[3].text = defense_start_time
    doc.paragraphs[6].runs[5].text = defense_end_time
    doc.paragraphs[17].runs[0].text = dit_head_full_name

    panel_table.cell(0, 9).paragraphs[0].runs[0].text = panel_response_date

    if panel_response == "accepted":
        panel_table.cell(0, 1).paragraphs[0].runs[0].text = "✓"
        panel_table.cell(0, 3).paragraphs[0].runs[0].text = "__"

    if panel_response == "declined":
        panel_table.cell(0, 1).paragraphs[0].runs[0].text = "__"
        panel_table.cell(0, 3).paragraphs[0].runs[0].text = "✓"


    if get_panel_invitation.dit_head_signature == 1:
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_invitation.dit_head_username) + ".png"):
            head_signature_table.cell(0, 0).text = ''
            head_signature = head_signature_table.cell(0, 0).add_paragraph()
            head_signature_run = head_signature.add_run()
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_invitation.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            get_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username)
            get_accepted_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
            get_pending_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_invitations": get_panel_invitations,
                "accepted_panel_invitations": get_accepted_panel_invitations.count(),
                "pending_panel_invitations": get_pending_panel_invitations.count(),
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-proposal-defense-panel-invitation.html", context)

    if get_panel_invitation.panel_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_invitation.panel_username) + ".png"):
            panel_signature_box.cell(0, 0).text = ''
            panel_signature = panel_signature_box.cell(0, 0).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_invitation.panel_username +'.png',width=Inches(1.2), height=Inches(0.45))


        else:
            get_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username)
            get_accepted_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
            get_pending_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_invitations": get_panel_invitations,
                "accepted_panel_invitations": get_accepted_panel_invitations.count(),
                "pending_panel_invitations": get_pending_panel_invitations.count(),
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet3-proposal-defense-panel-invitation.html", context)


    # Create - QR Code
    auth_qr_code = qrcode.make('Proposal Defense Panel Invitation\nH:' + dit_head_response_date + '\nP: ' + panel_response_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-" + panel_username + "-" + panel_response + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.docx")
    convert(current_user.username + "-" + panel_username + "-" + panel_response + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-" + panel_username + "-" + panel_response + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-" + panel_username + "-" + panel_response + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION-QR.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Panel Invitation for Proposal Defense QR - Deleted")
    else:
        print("Panel Invitation for Proposal Defense QR - Not Found")

    bet3_proposal_defense_panel_inviation_docx = current_user.username + "-" + panel_username + "-" + panel_response + "-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.docx"
    # bet3_proposal_defense_panel_inviation_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.docx')

    if os.path.isfile(bet3_proposal_defense_panel_inviation_docx):
        os.remove(bet3_proposal_defense_panel_inviation_docx)
        print("Panel Invitation for Proposal Defense - Deleted")
    else:
         print("Panel Invitation for Proposal Defense - Not Found")

    get_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username)
    get_accepted_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
    get_pending_panel_invitations = ProposalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        #'download_link': download_link,
        "panel_invitations": get_panel_invitations,
        "accepted_panel_invitations": get_accepted_panel_invitations.count(),
        "pending_panel_invitations": get_pending_panel_invitations.count(),
        "response": "sweet downloaded",
    }

    return render(request, "student-bet3-proposal-defense-panel-invitation-dashboard.html", context)


# Student - BET3 - Critique Form - Download
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentBET3CritiqueFormDownload(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    # Student - Get Student Leader Data
    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
    except:
        return redirect("login")
    
    # Get Student Group Members
    try:
        get_student_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
    except:
        get_student_group_members = None

    # Get Student Proposal Defense Accepted with Revision
    try:
        get_accepted_proposal_title = ResearchTitle.objects.get(student_leader_username=current_user.username, proposal_defense_status = "Accepted with Revision")
    except:
        print("pass research titles")
        return redirect("student-dashboard")

    if get_student_leader_data.group_members_status != "completed" and get_student_leader_data.research_titles_status != "completed" and get_student_leader_data.bet3_panel_invitation_status != "completed":
        return redirect("student-dashboard")

    if get_student_leader_data.middle_name == "":
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name
    else:
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name + " " + get_student_leader_data.middle_name[0] + "."

    get_critique_form = ProposalDefenseCritique.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_proposal_defense_date)
    get_proposal_defense_form = ProposalDefenseForm.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_proposal_defense_date)

    get_critique_form_panel_data = ProposalDefenseCritique.objects.all().filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = False, defense_date = get_student_leader_data.research_proposal_defense_date)
    get_critique_form_panel_chair_data = ProposalDefenseCritique.objects.filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = True, defense_date = get_student_leader_data.research_proposal_defense_date)

    doc = Document("static/forms/5-CRITIQUE-FORM.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/5-CRITIQUE-FORM.docx')
    
    paragraph = doc.add_paragraph()

    for i in range(len(get_critique_form)):
        if get_critique_form[i].critique:
            paragraph.add_run("     ● " + get_critique_form[i].critique)
            run = paragraph.add_run()
            run.add_break()
        i + 1
    

    panel_1 = doc.tables[1]
    panel_chair = doc.tables[2]
    panel_3 = doc.tables[3]
    panel_4 = doc.tables[5]
    panel_5 = doc.tables[6]
    panel_table_1 = doc.tables[4]
    panel_table_2 = doc.tables[8]
    date_table = doc.tables[7]
    qr_code_box = doc.tables[9]

    # Date of the Form
    date_table.cell(0, 1).text = get_critique_form_panel_chair_data[0].defense_date

    # Panel Chairman
    if get_critique_form_panel_chair_data[0]:
        panel_table_1.cell(1, 2).paragraphs[0].runs[0].text = get_critique_form_panel_chair_data[0].panel_full_name
        if get_critique_form_panel_chair_data[0].panel_chairman_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_chair_data[0].panel_username) + ".png"):
                panel_table_1.cell(0, 2).text = ""
                panel_chair_sign = panel_chair.cell(0, 0).add_paragraph()
                panel_chair_sign_run = panel_chair_sign.add_run()
                panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)

    # Panel 1
    if get_critique_form_panel_data[0]:
        panel_table_1.cell(1, 0).paragraphs[0].runs[0].text = get_critique_form_panel_data[0].panel_full_name
        if get_critique_form_panel_data[0].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[0].panel_username) + ".png"):
                panel_table_1.cell(0, 0).text = ""
                panel_sign = panel_1.cell(0, 0).add_paragraph()
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
    
    # Panel 2
    if get_critique_form_panel_data[1]:
        panel_table_1.cell(1, 4).paragraphs[0].runs[0].text = get_critique_form_panel_data[1].panel_full_name
        if get_critique_form_panel_data[1].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[1].panel_username) + ".png"):
                panel_table_1.cell(0, 4).text = ""
                panel_sign = panel_3.cell(0, 0).add_paragraph()
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[1].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
    
    # Panel 3
    try:
        if get_critique_form_panel_data[2]:
            panel_table_2.cell(1, 0).paragraphs[0].runs[0].text = get_critique_form_panel_data[2].panel_full_name
            if get_critique_form_panel_data[2].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[2].panel_username) + ".png"):
                    panel_table_2.cell(0, 0).text = ""
                    panel_sign = panel_4.cell(0, 0).add_paragraph()
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[2].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                    # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-bet3-critique-form.html", context)
    except:
        panel_table_2.cell(1, 0).paragraphs[0].runs[0].text = ""
        panel_table_2.cell(0, 0).text = ""

    try:
        # Panel 4
        if get_critique_form_panel_data[3]:
            panel_table_2.cell(1, 4).paragraphs[0].runs[0].text = get_critique_form_panel_data[3].panel_full_name
            if get_critique_form_panel_data[1].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[3].panel_username) + ".png"):
                    panel_table_2.cell(0, 4).text = ""
                    panel_sign = panel_5.cell(0, 0).add_paragraph()
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[3].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                    # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-bet3-critique-form.html", context)
    except:
        panel_table_2.cell(1, 4).paragraphs[0].runs[0].text = ""
        panel_table_2.cell(0, 4).text = ""


    # Create - QR Code
    auth_qr_code = qrcode.make('Critique Form\nDate:' + get_student_leader_data.research_proposal_defense_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-BET3-CRITIQUE-FORM-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET3-CRITIQUE-FORM-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-BET3-CRITIQUE-FORM.docx")
    convert(current_user.username + "-BET3-CRITIQUE-FORM.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-BET3-CRITIQUE-FORM.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-BET3-CRITIQUE-FORM.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-CRITIQUE-FORM.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-BET3-CRITIQUE-FORM.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-CRITIQUE-FORM.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-CRITIQUE-FORM.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET3-CRITIQUE-FORM-QR.png"

    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Critique Form QR - Deleted")
    else:
        print("Critique Form QR - Not Found")

    delete_critique_form = current_user.username + "-BET3-CRITIQUE-FORM.docx"
    # delete_critique_form = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-CRITIQUE-FORM.docx')

    if os.path.isfile(delete_critique_form):
        os.remove(delete_critique_form)
        print("Critique Form - Deleted")
    else:
        print("Critique Form - Form")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        "student_leader_full_name": student_leader_full_name,
        "student_group_members": get_student_group_members,
        "student_research_title": get_accepted_proposal_title,
        "critiques": get_critique_form,
        "proposal_defense_form": get_proposal_defense_form,
        "response": "sweet downloaded"
    }

    return render(request, "student-bet3-critique-form.html", context)


# Student - BET3 - Critique Form - Download
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentBET3ResearchProposalDefenseFormDownload(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    # Student - Get Student Leader Data
    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
    except:
        return redirect("login")

    # Get Student Group Members
    try:
        get_student_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
    except:
        get_student_group_members = None

    # Get Student Proposal Defense Accepted with Revision
    try:
        get_accepted_proposal_title = ResearchTitle.objects.get(student_leader_username=current_user.username, proposal_defense_status = "Accepted with Revision")
    except:
        print("pass research titles")
        return redirect("student-dashboard")

    if get_student_leader_data.group_members_status != "completed" and get_student_leader_data.research_titles_status != "completed" and get_student_leader_data.bet3_panel_invitation_status != "completed":
        return redirect("student-dashboard")

    if get_student_leader_data.middle_name == "":
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name
    else:
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name + " " + get_student_leader_data.middle_name[0] + "."

    get_critique_form = ProposalDefenseCritique.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_proposal_defense_date)
    get_proposal_defense_form = ProposalDefenseForm.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_proposal_defense_date)
    get_panel_chair_proposal_defense_form = ProposalDefenseForm.objects.filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_proposal_defense_date, is_panel_chairman = True)

    get_critique_form_panel_data = ProposalDefenseCritique.objects.all().filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = False, defense_date = get_student_leader_data.research_proposal_defense_date)
    get_critique_form_panel_chair_data = ProposalDefenseCritique.objects.filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = True, defense_date = get_student_leader_data.research_proposal_defense_date)

    get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)

    doc = Document("static/forms/6-RESEARCH-PROPOSAL-DEFENSE-FORM.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/6-RESEARCH-PROPOSAL-DEFENSE-FORM.docx')
    
    paragraph = doc.add_paragraph()
   
    examinee_table = doc.tables[1]
    panel_table = doc.tables[2]
    panelchair_table = doc.tables[3]
    qrcodebox_table = doc.tables[4]

    # Student Leader
    examinee_table.cell(0, 2).text = get_proposal_defense_form[0].student_leader_full_name

    # Student 2
    try:
        if get_group_members[0]:
            examinee_table.cell(1, 2).text = get_group_members[0].student_member_full_name
    except:
        examinee_table.cell(1, 2).text = ""


    # Student 3
    try:
        if get_group_members[1]:
            examinee_table.cell(2, 2).text = get_group_members[1].student_member_full_name
    except:
        examinee_table.cell(2, 2).text = ""

    # Student 4
    try:
        if get_group_members[2]:
            examinee_table.cell(3, 2).text = get_group_members[2].student_member_full_name
    except:
        examinee_table.cell(3, 2).text = ""

    # Student 5
    try:
        if get_group_members[3]:
            examinee_table.cell(4, 2).text = get_group_members[3].student_member_full_name
    except:
         examinee_table.cell(4, 2).text = ""
    
    # Degree
    examinee_table.cell(6, 2).text = get_student_leader_data.course_major_abbr

    # Title of Research
    examinee_table.cell(8, 2).text = get_accepted_proposal_title.research_title

    # Defense Date
    examinee_table.cell(9, 2).text = get_proposal_defense_form[0].defense_date

    # Defense Time
    examinee_table.cell(10, 2).text = get_proposal_defense_form[0].defense_start_time + " - " + get_proposal_defense_form[0].defense_end_time

    # Thesis Adviser name
    examinee_table.cell(12, 2).text = get_student_leader_data.adviser_name
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Panel 1
    if get_proposal_defense_form[0]:
        panel_table.cell(1, 0).text = ""
        panel_name = panel_table.cell(1, 0).add_paragraph(get_proposal_defense_form[0].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_proposal_defense_form[0].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[0].panel_username) + ".png"):
                if get_proposal_defense_form[0].proposal_defense_response == "Accepted with Revision":
                    panel_table.cell(1, 1).text = ""
                    panel_table.cell(1, 2).text = ""
                    panel_table.cell(1, 3).text = ""
                    panel_sign = panel_table.cell(1, 1).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[0].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                
                elif get_proposal_defense_form[0].proposal_defense_response == "Deferred with Revision":
                    panel_table.cell(1, 1).text = ""
                    panel_table.cell(1, 2).text = ""
                    panel_table.cell(1, 3).text = ""
                    panel_sign = panel_table.cell(1, 2).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[0].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                elif get_proposal_defense_form[0].proposal_defense_response == "Not Accepted":
                    panel_table.cell(1, 1).text = ""
                    panel_table.cell(1, 2).text = ""
                    panel_table.cell(1, 3).text = ""
                    panel_sign = panel_table.cell(1, 3).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[0].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
        else:
            if get_proposal_defense_form[0].proposal_defense_response == "Accepted with Revision":
                panel_table.cell(1, 2).text = ""
                panel_table.cell(1, 3).text = ""
                
            elif get_proposal_defense_form[0].proposal_defense_response == "Deferred with Revision":
                panel_table.cell(1, 1).text = ""
                panel_table.cell(1, 3).text = "" 

            elif get_proposal_defense_form[0].proposal_defense_response == "Not Accepted":
                panel_table.cell(1, 1).text = ""
                panel_table.cell(1, 2).text = ""   


    # Panel 2
    if get_proposal_defense_form[1]:
        panel_table.cell(2, 0).text = ""
        panel_name = panel_table.cell(2, 0).add_paragraph(get_proposal_defense_form[1].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_proposal_defense_form[1].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[1].panel_username) + ".png"):
                if get_proposal_defense_form[1].proposal_defense_response == "Accepted with Revision":
                    panel_table.cell(2, 1).text = ""
                    panel_table.cell(2, 2).text = ""
                    panel_table.cell(2, 3).text = ""
                    panel_sign = panel_table.cell(2, 1).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[1].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                
                elif get_proposal_defense_form[1].proposal_defense_response == "Deferred with Revision":
                    panel_table.cell(2, 1).text = ""
                    panel_table.cell(2, 2).text = ""
                    panel_table.cell(2, 3).text = ""
                    panel_sign = panel_table.cell(2, 2).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[1].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                elif get_proposal_defense_form[1].proposal_defense_response == "Not Accepted":
                    panel_table.cell(2, 1).text = ""
                    panel_table.cell(2, 2).text = ""
                    panel_table.cell(2, 3).text = ""
                    panel_sign = panel_table.cell(2, 3).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[1].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    

            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
        else:
            if get_proposal_defense_form[1].proposal_defense_response == "Accepted with Revision":
                panel_table.cell(2, 2).text = ""
                panel_table.cell(2, 3).text = ""
                
            elif get_proposal_defense_form[1].proposal_defense_response == "Deferred with Revision":
                panel_table.cell(2, 1).text = ""
                panel_table.cell(2, 3).text = "" 

            elif get_proposal_defense_form[1].proposal_defense_response == "Not Accepted":
                panel_table.cell(2, 1).text = ""
                panel_table.cell(2, 2).text = ""    


    # Panel 3
    if get_proposal_defense_form[2]:
        panel_table.cell(3, 0).text = ""
        panel_name = panel_table.cell(3, 0).add_paragraph(get_proposal_defense_form[2].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_proposal_defense_form[2].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[2].panel_username) + ".png"):
                if get_proposal_defense_form[2].proposal_defense_response == "Accepted with Revision":
                    panel_table.cell(3, 1).text = ""
                    panel_table.cell(3, 2).text = ""
                    panel_table.cell(3, 3).text = ""
                    panel_sign = panel_table.cell(3, 1).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[2].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                
                elif get_proposal_defense_form[2].proposal_defense_response == "Deferred with Revision":
                    panel_table.cell(3, 1).text = ""
                    panel_table.cell(3, 2).text = ""
                    panel_table.cell(3, 3).text = ""
                    panel_sign = panel_table.cell(3, 2).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[2].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                elif get_proposal_defense_form[2].proposal_defense_response == "Not Accepted":
                    panel_table.cell(3, 1).text = ""
                    panel_table.cell(3, 2).text = ""
                    panel_table.cell(3, 3).text = ""
                    panel_sign = panel_table.cell(3, 3).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[2].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    

            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
        else:
            if get_proposal_defense_form[2].proposal_defense_response == "Accepted with Revision":
                panel_table.cell(3, 2).text = ""
                panel_table.cell(3, 3).text = ""
                
            elif get_proposal_defense_form[2].proposal_defense_response == "Deferred with Revision":
                panel_table.cell(3, 1).text = ""
                panel_table.cell(3, 3).text = "" 

            elif get_proposal_defense_form[2].proposal_defense_response == "Not Accepted":
                panel_table.cell(3, 1).text = ""
                panel_table.cell(3, 2).text = ""


    # Panel 4
    try:
        if get_proposal_defense_form[3]:
            panel_table.cell(4, 0).text = ""
            panel_name = panel_table.cell(4, 0).add_paragraph(get_proposal_defense_form[3].panel_full_name)
            panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if get_proposal_defense_form[3].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[3].panel_username) + ".png"):
                    if get_proposal_defense_form[3].proposal_defense_response == "Accepted with Revision":
                        panel_table.cell(4, 1).text = ""
                        panel_table.cell(4, 2).text = ""
                        panel_table.cell(4, 3).text = ""
                        panel_sign = panel_table.cell(4, 1).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[3].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    
                    elif get_proposal_defense_form[3].proposal_defense_response == "Deferred with Revision":
                        panel_table.cell(4, 1).text = ""
                        panel_table.cell(4, 2).text = ""
                        panel_table.cell(4, 3).text = ""
                        panel_sign = panel_table.cell(4, 2).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[3].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                    elif get_proposal_defense_form[3].proposal_defense_response == "Not Accepted":
                        panel_table.cell(4, 1).text = ""
                        panel_table.cell(4, 2).text = ""
                        panel_table.cell(4, 3).text = ""
                        panel_sign = panel_table.cell(4, 3).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[3].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                        

                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-bet3-critique-form.html", context)
            else:
                if get_proposal_defense_form[3].proposal_defense_response == "Accepted with Revision":
                    panel_table.cell(4, 2).text = ""
                    panel_table.cell(4, 3).text = ""
                    
                elif get_proposal_defense_form[3].proposal_defense_response == "Deferred with Revision":
                    panel_table.cell(4, 1).text = ""
                    panel_table.cell(4, 3).text = "" 

                elif get_proposal_defense_form[3].proposal_defense_response == "Not Accepted":
                    panel_table.cell(4, 1).text = ""
                    panel_table.cell(4, 2).text = ""        
    
    except:
        panel_table.cell(4, 0).text = ""
        panel_table.cell(4, 1).text = ""
        panel_table.cell(4, 2).text = "" 
        panel_table.cell(4, 3).text = "" 

    # Panel 5
    try:
        if get_proposal_defense_form[4]:
            panel_table.cell(5, 0).text = ""
            panel_name = panel_table.cell(5, 0).add_paragraph(get_proposal_defense_form[4].panel_full_name)
            panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if get_proposal_defense_form[4].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[4].panel_username) + ".png"):
                    if get_proposal_defense_form[4].proposal_defense_response == "Accepted with Revision":
                        panel_table.cell(5, 1).text = ""
                        panel_table.cell(5, 2).text = ""
                        panel_table.cell(5, 3).text = ""
                        panel_sign = panel_table.cell(5, 1).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[4].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    
                    elif get_proposal_defense_form[4].proposal_defense_response == "Deferred with Revision":
                        panel_table.cell(5, 1).text = ""
                        panel_table.cell(5, 2).text = ""
                        panel_table.cell(5, 3).text = ""
                        panel_sign = panel_table.cell(5, 2).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[4].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                    elif get_proposal_defense_form[4].proposal_defense_response == "Not Accepted":
                        panel_table.cell(5, 1).text = ""
                        panel_table.cell(5, 2).text = ""
                        panel_table.cell(5, 3).text = ""
                        panel_sign = panel_table.cell(5, 3).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[4].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                        

                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-bet3-critique-form.html", context)
            else:
                if get_proposal_defense_form[4].proposal_defense_response == "Accepted with Revision":
                    panel_table.cell(5, 2).text = ""
                    panel_table.cell(5, 3).text = ""
                    
                elif get_proposal_defense_form[4].proposal_defense_response == "Deferred with Revision":
                    panel_table.cell(5, 1).text = ""
                    panel_table.cell(5, 3).text = "" 

                elif get_proposal_defense_form[4].proposal_defense_response == "Not Accepted":
                    panel_table.cell(5, 1).text = ""
                    panel_table.cell(5, 2).text = ""  
    
    except:
        panel_table.cell(5, 0).text = ""
        panel_table.cell(5, 1).text = ""
        panel_table.cell(5, 2).text = ""
        panel_table.cell(5, 3).text = ""
    

    # Panel 1
    if get_panel_chair_proposal_defense_form[0]:
        panelchair_table.cell(1, 0).text = ""
        panel_name = panelchair_table.cell(1, 0).add_paragraph(get_panel_chair_proposal_defense_form[0].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_panel_chair_proposal_defense_form[0].panel_chairman_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_chair_proposal_defense_form[0].panel_username) + ".png"):
                panelchair_table.cell(0, 0).text = ""
                panel_sign = panelchair_table.cell(0, 0).add_paragraph()
                panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_chair_proposal_defense_form[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)

    # Create - QR Code
    auth_qr_code = qrcode.make('Research Proposal Defense Form\nDate:' + get_student_leader_data.research_proposal_defense_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM-QR.png")

    # INSERT IMAGE
    qr_code = qrcodebox_table.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.docx")
    convert(current_user.username + "-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM-QR.png"

    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Research Proposal Defense Form QR - Deleted")
    else:
        print("Research Proposal Defense Form QR - Not Found")

    delete_critique_form = current_user.username+"-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM-QR.docx"
    #delete_critique_form = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/"+'-BET3-RESEARCH-PROPOSAL-DEFENSE-FORM.docx')

    if os.path.isfile(delete_critique_form):
        os.remove(delete_critique_form)
        print("Research Proposal Defense Form - Deleted")
    else:
        print("Research Proposal Defense Form - Not Found")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        "student_leader_full_name": student_leader_full_name,
        "student_group_members": get_student_group_members,
        "student_research_title": get_accepted_proposal_title,
        "critiques": get_critique_form,
        "proposal_defense_form": get_proposal_defense_form,
        "response": "sweet downloaded"
    }

    return render(request, "student-bet3-research-proposal-defense.html", context)


# Student - BET5 - Final Defense - Panel Invitation - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentBET5FinalDefensePanelInvitationDownload(request, id):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
        get_panel_invitation = FinalPanelInvitation.objects.get(id=int(id), student_leader_username=current_user.username)

    except:
        return redirect("student-bet5-final-defense-panel-invitation-dashboard")
    
    try:
        get_accepted_research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, title_defense_status = "Accepted")
        research_title = get_accepted_research_title.research_title
    except:
        pass

    try:
        get_revise_research_title = ResearchTitle.objects.get(student_leader_username=current_user.username, title_defense_status = "Revise Title")
        research_title = get_revise_research_title.research_title
    except:
        pass


    ############## BET-3 PANEL INVITATION DATA ##############
    date_submitted = get_panel_invitation.form_date_sent

    student_member_list = [get_panel_invitation.student_leader_full_name]
    course = get_student_leader_data.course.replace("Engineering", "Eng.")
    major = get_student_leader_data.major


    defense_date = get_panel_invitation.research_final_defense_date
    defense_start_time = get_panel_invitation.research_final_defense_start_time
    defense_end_time = get_panel_invitation.research_final_defense_end_time

    dit_head_full_name = get_panel_invitation.dit_head_full_name
    dit_head_response = get_panel_invitation.dit_head_response
    dit_head_response_date = get_panel_invitation.dit_head_response_date

    panel_full_name = get_panel_invitation.panel_full_name
    panel_username = get_panel_invitation.panel_username
    panel_response = get_panel_invitation.panel_response
    panel_response_date = get_panel_invitation.panel_response_date
    ############## BET-3 PANEL INVITATION DATA ##############

    if get_group_members:
        for group_member in get_group_members:
            student_member_list.append(group_member.student_member_full_name)

    student_member_list.sort()

    print("Date Sent: ", date_submitted)

    print("Group Members: ", student_member_list)
    print("Course: ", course)
    print("Major: ", major)
    print("Research Title: ", research_title)

    print("Defense Date: ", defense_date)
    print("Defense Start Time: ", defense_start_time)
    print("Defense End Time: ", defense_end_time)

    print("DIT Head Name: ", dit_head_full_name)
    print("DIT Head Response: ", dit_head_response)
    print("DIT Head Response Date: ", dit_head_response_date)

    print("Panel Name: ", panel_full_name)
    print("Panel Response: ", panel_response)
    print("Panel Response Date: ", panel_response_date)

    doc = Document("static/forms/7-FINAL-DEFENSE-PANEL-INVITATION.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/7-FINAL-DEFENSE-PANEL-INVITATION.docx')

    student_table = doc.tables[1] # Student Data Table
    head_signature_table = doc.tables[2]
    panel_table = doc.tables[4] # Panel Data Table
    qr_code_box = doc.tables[5]
    panel_signature_box = doc.tables[3]

    try:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = student_member_list[0]
        student_table.cell(1, 2).paragraphs[0].runs[0].text = course
        student_table.cell(1, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(1, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(1, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = student_member_list[1]
        student_table.cell(2, 2).paragraphs[0].runs[0].text = course
        student_table.cell(2, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(2, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(2, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = student_member_list[2]
        student_table.cell(3, 2).paragraphs[0].runs[0].text = course
        student_table.cell(3, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(3, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(3, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = student_member_list[3]
        student_table.cell(4, 2).paragraphs[0].runs[0].text = course
        student_table.cell(4, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(4, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(4, 4).paragraphs[0].runs[0].text = ""

    try:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = student_member_list[4]
        student_table.cell(5, 2).paragraphs[0].runs[0].text = course
        student_table.cell(5, 4).paragraphs[0].runs[0].text = major
    except:
        student_table.cell(5, 0).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 2).paragraphs[0].runs[0].text = ""
        student_table.cell(5, 4).paragraphs[0].runs[0].text = ""

    doc.paragraphs[0].runs[1].text = date_submitted
    doc.paragraphs[1].runs[1].text = panel_full_name

    doc.paragraphs[4].runs[5].text = research_title

    doc.paragraphs[6].runs[3].text = defense_date
    doc.paragraphs[6].runs[5].text = defense_start_time
    doc.paragraphs[6].runs[9].text = defense_end_time
    doc.paragraphs[17].runs[0].text = dit_head_full_name

    panel_table.cell(0, 9).paragraphs[0].runs[0].text = panel_response_date

    if panel_response == "accepted":
        panel_table.cell(0, 1).paragraphs[0].runs[0].text = "✓"
        panel_table.cell(0, 3).paragraphs[0].runs[0].text = "__"

    if panel_response == "declined":
        panel_table.cell(0, 1).paragraphs[0].runs[0].text = "__"
        panel_table.cell(0, 3).paragraphs[0].runs[0].text = "✓"


    if get_panel_invitation.dit_head_signature == 1:
        # Check if DIT Head E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_invitation.dit_head_username) + ".png"):
            head_signature_table.cell(0, 0).text = ''
            head_signature = head_signature_table.cell(0, 0).add_paragraph()
            head_signature_run = head_signature.add_run()
            head_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_invitation.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            get_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username)
            get_accepted_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
            get_pending_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_invitations": get_panel_invitations,
                "accepted_panel_invitations": get_accepted_panel_invitations.count(),
                "pending_panel_invitations": get_pending_panel_invitations.count(),
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet5-final-defense-panel-invitation.html", context)

    if get_panel_invitation.panel_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_invitation.panel_username) + ".png"):
            panel_signature_box.cell(0, 0).text = ''
            panel_signature = panel_signature_box.cell(0, 0).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_invitation.panel_username +'.png',width=Inches(1.2), height=Inches(0.45))


        else:
            get_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username)
            get_accepted_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
            get_pending_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "panel_invitations": get_panel_invitations,
                "accepted_panel_invitations": get_accepted_panel_invitations.count(),
                "pending_panel_invitations": get_pending_panel_invitations.count(),
                "response": "sweet faculty member no signature",
            }

            return render(request, "student-bet5-final-defense-panel-invitation.html", context)


    # Create - QR Code
    auth_qr_code = qrcode.make('Final Defense Panel Invitation\nH:' + dit_head_response_date + '\nP: ' + panel_response_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-BET5-FINAL-DEFENSE-PANEL-INVITATION-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET5-FINAL-DEFENSE-PANEL-INVITATION-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-" + panel_username + "-" + panel_response + "-BET5-FINAL-DEFENSE-PANEL-INVITATION.docx")
    convert(current_user.username + "-" + panel_username + "-" + panel_response + "-BET5-FINAL-DEFENSE-PANEL-INVITATION.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-" + panel_username + "-" + panel_response + "-BET5-FINAL-DEFENSE-PANEL-INVITATION.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-" + panel_username + "-" + panel_response + "-BET5-FINAL-DEFENSE-PANEL-INVITATION.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-BET5-FINAL-DEFENSE-PANEL-INVITATION.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-BET5-FINAL-DEFENSE-PANEL-INVITATION.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-BET5-FINAL-DEFENSE-PANEL-INVITATION.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-BET5-FINAL-DEFENSE-PANEL-INVITATION.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET5-FINAL-DEFENSE-PANEL-INVITATION-QR.png"
    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Panel Invitation for Final Defense QR - Deleted")
    else:
        print("Panel Invitation for Final Defense QR - Not Found")

    bet3_proposal_defense_panel_inviation_docx = current_user.username + "-" + panel_username + "-" + panel_response + "-BET5-FINAL-DEFENSE-PANEL-INVITATION.docx"
    # bet3_proposal_defense_panel_inviation_docx = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-BET5-FINAL-DEFENSE-PANEL-INVITATION.docx')

    if os.path.isfile(bet3_proposal_defense_panel_inviation_docx):
        os.remove(bet3_proposal_defense_panel_inviation_docx)
        print("Panel Invitation for Final Defense - Deleted")
    else:
        print("Panel Invitation for Final Defense - Not Found")

    get_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username)
    get_accepted_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="accepted")
    get_pending_panel_invitations = FinalPanelInvitation.objects.all().filter(student_leader_username=current_user.username, form_status="pending")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        #'download_link': download_link,
        "panel_invitations": get_panel_invitations,
        "accepted_panel_invitations": get_accepted_panel_invitations.count(),
        "pending_panel_invitations": get_pending_panel_invitations.count(),
        "response": "sweet downloaded",
    }

    return render(request, "student-bet5-final-defense-panel-invitation-dashboard.html", context)


# Student - BET3 - Critique Form - Download
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentFinalCritiqueFormDownload(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    # Student - Get Student Leader Data
    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
    except:
        return redirect("login")
    
    # Get Student Group Members
    try:
        get_student_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
    except:
        get_student_group_members = None

    # Get Student Proposal Defense Accepted with Revision
    try:
        get_accepted_proposal_title = ResearchTitle.objects.get(student_leader_username=current_user.username, proposal_defense_status = "Accepted with Revision")
    except:
        print("pass research titles")
        return redirect("student-dashboard")

    if get_student_leader_data.group_members_status != "completed" and get_student_leader_data.research_titles_status != "completed" and get_student_leader_data.bet3_panel_invitation_status != "completed":
        return redirect("student-dashboard")

    if get_student_leader_data.middle_name == "":
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name
    else:
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name + " " + get_student_leader_data.middle_name[0] + "."

    get_critique_form = FinalDefenseCritique.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_final_defense_date)
    get_proposal_defense_form = FinalDefenseForm.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_final_defense_date)

    get_critique_form_panel_data = FinalDefenseCritique.objects.all().filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = False, defense_date = get_student_leader_data.research_final_defense_date)
    get_critique_form_panel_chair_data = FinalDefenseCritique.objects.filter(student_leader_username=current_user.username, critique = "", is_panel_chairman = True, defense_date = get_student_leader_data.research_final_defense_date)

    doc = Document("static/forms/5-CRITIQUE-FORM.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/5-CRITIQUE-FORM.docx')
    
    paragraph = doc.add_paragraph()

    for i in range(len(get_critique_form)):
        if get_critique_form[i].critique:
            paragraph.add_run("     ● " + get_critique_form[i].critique)
            run = paragraph.add_run()
            run.add_break()
        i + 1
    

    panel_1 = doc.tables[1]
    panel_chair = doc.tables[2]
    panel_3 = doc.tables[3]
    panel_4 = doc.tables[5]
    panel_5 = doc.tables[6]
    panel_table_1 = doc.tables[4]
    panel_table_2 = doc.tables[8]
    date_table = doc.tables[7]
    qr_code_box = doc.tables[9]

    # Date of the Form
    date_table.cell(0, 1).text = get_critique_form_panel_chair_data[0].defense_date

    # Panel Chairman
    if get_critique_form_panel_chair_data[0]:
        panel_table_1.cell(1, 2).paragraphs[0].runs[0].text = get_critique_form_panel_chair_data[0].panel_full_name
        if get_critique_form_panel_chair_data[0].panel_chairman_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_chair_data[0].panel_username) + ".png"):
                panel_table_1.cell(0, 2).text = ""
                panel_chair_sign = panel_chair.cell(0, 0).add_paragraph()
                panel_chair_sign_run = panel_chair_sign.add_run()
                panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-final-critique-form.html", context)

    # Panel 1
    if get_critique_form_panel_data[0]:
        panel_table_1.cell(1, 0).paragraphs[0].runs[0].text = get_critique_form_panel_data[0].panel_full_name
        if get_critique_form_panel_data[0].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[0].panel_username) + ".png"):
                panel_table_1.cell(0, 0).text = ""
                panel_sign = panel_1.cell(0, 0).add_paragraph()
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-final-critique-form.html", context)
    
    # Panel 2
    if get_critique_form_panel_data[1]:
        panel_table_1.cell(1, 4).paragraphs[0].runs[0].text = get_critique_form_panel_data[1].panel_full_name
        if get_critique_form_panel_data[1].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[1].panel_username) + ".png"):
                panel_table_1.cell(0, 4).text = ""
                panel_sign = panel_3.cell(0, 0).add_paragraph()
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[1].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "critiques": get_critique_form,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-final-critique-form.html", context)
    
    # Panel 3
    try:
        if get_critique_form_panel_data[2]:
            panel_table_2.cell(1, 0).paragraphs[0].runs[0].text = get_critique_form_panel_data[2].panel_full_name
            if get_critique_form_panel_data[2].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[2].panel_username) + ".png"):
                    panel_table_2.cell(0, 0).text = ""
                    panel_sign = panel_4.cell(0, 0).add_paragraph()
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[2].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                    # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-final-critique-form.html", context)
    except:
        panel_table_2.cell(1, 0).paragraphs[0].runs[0].text = ""
        panel_table_2.cell(0, 0).text = ""

    try:
        # Panel 4
        if get_critique_form_panel_data[3]:
            panel_table_2.cell(1, 4).paragraphs[0].runs[0].text = get_critique_form_panel_data[3].panel_full_name
            if get_critique_form_panel_data[1].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_critique_form_panel_data[3].panel_username) + ".png"):
                    panel_table_2.cell(0, 4).text = ""
                    panel_sign = panel_5.cell(0, 0).add_paragraph()
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_data[3].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                    # panel_chair_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_critique_form_panel_chair_data[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "critiques": get_critique_form,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-final-critique-form.html", context)
    except:
        panel_table_2.cell(1, 4).paragraphs[0].runs[0].text = ""
        panel_table_2.cell(0, 4).text = ""


    # Create - QR Code
    auth_qr_code = qrcode.make('Critique Form\nDate:' + get_student_leader_data.research_proposal_defense_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-BET3-CRITIQUE-FORM-QR.png")

    # INSERT IMAGE
    qr_code = qr_code_box.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET3-CRITIQUE-FORM-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-FINAL-CRITIQUE-FORM.docx")
    convert(current_user.username + "-FINAL-CRITIQUE-FORM.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-FINAL-CRITIQUE-FORM.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-FINAL-CRITIQUE-FORM.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-FINAL-CRITIQUE-FORM.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+"-"+panel_username+"-"+panel_response+'-FINAL-CRITIQUE-FORM.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username +"-"+panel_username+"-"+panel_response+'-FINAL-CRITIQUE-FORM.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-FINAL-CRITIQUE-FORM.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET3-CRITIQUE-FORM-QR.png"

    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Critique Form QR - Deleted")
    else:
        print("Critique Form QR - Not Found")

    delete_critique_form = current_user.username + "-FINAL-CRITIQUE-FORM.docx"
    # delete_critique_form = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-FINAL-CRITIQUE-FORM.docx')

    if os.path.isfile(delete_critique_form):
        os.remove(delete_critique_form)
        print("Critique Form - Deleted")
    else:
        print("Critique Form - Form")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        "student_leader_full_name": student_leader_full_name,
        "student_group_members": get_student_group_members,
        "student_research_title": get_accepted_proposal_title,
        "critiques": get_critique_form,
        "proposal_defense_form": get_proposal_defense_form,
        "response": "sweet downloaded"
    }

    return render(request, "student-final-critique-form.html", context)

# Student - BET3 - Critique Form - Download
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentBET5ResearchFinalDefenseFormDownload(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    # Student - Get Student Leader Data
    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
    except:
        return redirect("login")

    # Get Student Group Members
    try:
        get_student_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
    except:
        get_student_group_members = None

    # Get Student Proposal Defense Accepted with Revision
    try:
        get_accepted_proposal_title = ResearchTitle.objects.get(student_leader_username=current_user.username, final_defense_status = "Accepted with Revision")
    except:
        print("pass research titles")
        return redirect("student-dashboard")

    if get_student_leader_data.group_members_status != "completed" and get_student_leader_data.research_titles_status != "completed" and get_student_leader_data.bet5_final_defense_panel_invitation_status != "completed":
        return redirect("student-dashboard")

    if get_student_leader_data.middle_name == "":
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name
    else:
        student_leader_full_name = get_student_leader_data.last_name + " " + get_student_leader_data.suffix + ", " + get_student_leader_data.first_name + " " + get_student_leader_data.middle_name[0] + "."

    get_proposal_defense_form = FinalDefenseForm.objects.all().filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_final_defense_date)
    get_panel_chair_proposal_defense_form = FinalDefenseForm.objects.filter(student_leader_username=current_user.username, defense_date = get_student_leader_data.research_final_defense_date, is_panel_chairman = True)


    get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)

    doc = Document("static/forms/8-RESEARCH-FINAL-DEFENSE-FORM.docx")
    # doc = Document('/home/johnanthonybataller/tupc-research-defense-form-django/static/forms/8-RESEARCH-FINAL-DEFENSE-FORM.docx')
    
    paragraph = doc.add_paragraph()
   
    examinee_table = doc.tables[1]
    panel_table = doc.tables[2]
    panelchair_table = doc.tables[3]
    qrcodebox_table = doc.tables[4]

    # Student Leader
    examinee_table.cell(0, 2).text = get_proposal_defense_form[0].student_leader_full_name

    # Student 2
    try:
        if get_group_members[0]:
            examinee_table.cell(1, 2).text = get_group_members[0].student_member_full_name
    except:
        examinee_table.cell(1, 2).text = ""


    # Student 3
    try:
        if get_group_members[1]:
            examinee_table.cell(2, 2).text = get_group_members[1].student_member_full_name
    except:
        examinee_table.cell(2, 2).text = ""

    # Student 4
    try:
        if get_group_members[2]:
            examinee_table.cell(3, 2).text = get_group_members[2].student_member_full_name
    except:
        examinee_table.cell(3, 2).text = ""

    # Student 5
    try:
        if get_group_members[3]:
            examinee_table.cell(4, 2).text = get_group_members[3].student_member_full_name
    except:
         examinee_table.cell(4, 2).text = ""
    
    # Degree
    examinee_table.cell(6, 2).text = get_student_leader_data.course_major_abbr

    # Title of Research
    examinee_table.cell(8, 2).text = get_accepted_proposal_title.research_title

    # Defense Date
    examinee_table.cell(9, 2).text = get_proposal_defense_form[0].defense_date

    # Defense Time
    examinee_table.cell(10, 2).text = get_proposal_defense_form[0].defense_start_time + " - " + get_proposal_defense_form[0].defense_end_time

    # Thesis Adviser name
    examinee_table.cell(12, 2).text = get_student_leader_data.adviser_name
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Panel 1
    if get_proposal_defense_form[0]:
        panel_table.cell(1, 0).text = ""
        panel_name = panel_table.cell(1, 0).add_paragraph(get_proposal_defense_form[0].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_proposal_defense_form[0].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[0].panel_username) + ".png"):
                if get_proposal_defense_form[0].final_defense_response == "Accepted with Revision":
                    panel_table.cell(1, 1).text = ""
                    panel_table.cell(1, 2).text = ""
                    panel_table.cell(1, 3).text = ""
                    panel_sign = panel_table.cell(1, 1).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[0].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                
                elif get_proposal_defense_form[0].proposal_defense_response == "Deferred with Revision":
                    panel_table.cell(1, 1).text = ""
                    panel_table.cell(1, 2).text = ""
                    panel_table.cell(1, 3).text = ""
                    panel_sign = panel_table.cell(1, 2).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[0].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                elif get_proposal_defense_form[0].proposal_defense_response == "Not Accepted":
                    panel_table.cell(1, 1).text = ""
                    panel_table.cell(1, 2).text = ""
                    panel_table.cell(1, 3).text = ""
                    panel_sign = panel_table.cell(1, 3).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[0].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
        else:
            if get_proposal_defense_form[0].final_defense_response == "Accepted with Revision":
                panel_table.cell(1, 2).text = ""
                panel_table.cell(1, 3).text = ""
                
            elif get_proposal_defense_form[0].final_defense_response == "Deferred with Revision":
                panel_table.cell(1, 1).text = ""
                panel_table.cell(1, 3).text = "" 

            elif get_proposal_defense_form[0].final_defense_response == "Not Accepted":
                panel_table.cell(1, 1).text = ""
                panel_table.cell(1, 2).text = ""   


    # Panel 2
    if get_proposal_defense_form[1]:
        panel_table.cell(2, 0).text = ""
        panel_name = panel_table.cell(2, 0).add_paragraph(get_proposal_defense_form[1].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_proposal_defense_form[1].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[1].panel_username) + ".png"):
                if get_proposal_defense_form[1].final_defense_response == "Accepted with Revision":
                    panel_table.cell(2, 1).text = ""
                    panel_table.cell(2, 2).text = ""
                    panel_table.cell(2, 3).text = ""
                    panel_sign = panel_table.cell(2, 1).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[1].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                
                elif get_proposal_defense_form[1].final_defense_response == "Deferred with Revision":
                    panel_table.cell(2, 1).text = ""
                    panel_table.cell(2, 2).text = ""
                    panel_table.cell(2, 3).text = ""
                    panel_sign = panel_table.cell(2, 2).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[1].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                elif get_proposal_defense_form[1].final_defense_response == "Not Accepted":
                    panel_table.cell(2, 1).text = ""
                    panel_table.cell(2, 2).text = ""
                    panel_table.cell(2, 3).text = ""
                    panel_sign = panel_table.cell(2, 3).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[1].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    

            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
        else:
            if get_proposal_defense_form[1].final_defense_response == "Accepted with Revision":
                panel_table.cell(2, 2).text = ""
                panel_table.cell(2, 3).text = ""
                
            elif get_proposal_defense_form[1].final_defense_response == "Deferred with Revision":
                panel_table.cell(2, 1).text = ""
                panel_table.cell(2, 3).text = "" 

            elif get_proposal_defense_form[1].final_defense_response == "Not Accepted":
                panel_table.cell(2, 1).text = ""
                panel_table.cell(2, 2).text = ""    


    # Panel 3
    if get_proposal_defense_form[2]:
        panel_table.cell(3, 0).text = ""
        panel_name = panel_table.cell(3, 0).add_paragraph(get_proposal_defense_form[2].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_proposal_defense_form[2].panel_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[2].panel_username) + ".png"):
                if get_proposal_defense_form[2].final_defense_response == "Accepted with Revision":
                    panel_table.cell(3, 1).text = ""
                    panel_table.cell(3, 2).text = ""
                    panel_table.cell(3, 3).text = ""
                    panel_sign = panel_table.cell(3, 1).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[2].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                
                elif get_proposal_defense_form[2].final_defense_response == "Deferred with Revision":
                    panel_table.cell(3, 1).text = ""
                    panel_table.cell(3, 2).text = ""
                    panel_table.cell(3, 3).text = ""
                    panel_sign = panel_table.cell(3, 2).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[2].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                elif get_proposal_defense_form[2].final_defense_response == "Not Accepted":
                    panel_table.cell(3, 1).text = ""
                    panel_table.cell(3, 2).text = ""
                    panel_table.cell(3, 3).text = ""
                    panel_sign = panel_table.cell(3, 3).add_paragraph()
                    panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    panel_sign_run = panel_sign.add_run()
                    panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[2].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    

            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)
        else:
            if get_proposal_defense_form[2].final_defense_response == "Accepted with Revision":
                panel_table.cell(3, 2).text = ""
                panel_table.cell(3, 3).text = ""
                
            elif get_proposal_defense_form[2].final_defense_response == "Deferred with Revision":
                panel_table.cell(3, 1).text = ""
                panel_table.cell(3, 3).text = "" 

            elif get_proposal_defense_form[2].final_defense_response == "Not Accepted":
                panel_table.cell(3, 1).text = ""
                panel_table.cell(3, 2).text = ""


    # Panel 4
    try:
        if get_proposal_defense_form[3]:
            panel_table.cell(4, 0).text = ""
            panel_name = panel_table.cell(4, 0).add_paragraph(get_proposal_defense_form[3].panel_full_name)
            panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if get_proposal_defense_form[3].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[3].panel_username) + ".png"):
                    if get_proposal_defense_form[3].final_defense_response == "Accepted with Revision":
                        panel_table.cell(4, 1).text = ""
                        panel_table.cell(4, 2).text = ""
                        panel_table.cell(4, 3).text = ""
                        panel_sign = panel_table.cell(4, 1).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[3].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    
                    elif get_proposal_defense_form[3].final_defense_response == "Deferred with Revision":
                        panel_table.cell(4, 1).text = ""
                        panel_table.cell(4, 2).text = ""
                        panel_table.cell(4, 3).text = ""
                        panel_sign = panel_table.cell(4, 2).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[3].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                    elif get_proposal_defense_form[3].final_defense_response == "Not Accepted":
                        panel_table.cell(4, 1).text = ""
                        panel_table.cell(4, 2).text = ""
                        panel_table.cell(4, 3).text = ""
                        panel_sign = panel_table.cell(4, 3).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[3].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                        

                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-bet3-critique-form.html", context)
            else:
                if get_proposal_defense_form[3].final_defense_response == "Accepted with Revision":
                    panel_table.cell(4, 2).text = ""
                    panel_table.cell(4, 3).text = ""
                    
                elif get_proposal_defense_form[3].final_defense_response == "Deferred with Revision":
                    panel_table.cell(4, 1).text = ""
                    panel_table.cell(4, 3).text = "" 

                elif get_proposal_defense_form[3].final_defense_response == "Not Accepted":
                    panel_table.cell(4, 1).text = ""
                    panel_table.cell(4, 2).text = ""
    except:
        panel_table.cell(4, 0).text = ""
        panel_table.cell(4, 1).text = ""
        panel_table.cell(4, 2).text = ""
        panel_table.cell(4, 3).text = ""
    

    # Panel 5
    try:
        if get_proposal_defense_form[4]:
            panel_table.cell(5, 0).text = ""
            panel_name = panel_table.cell(5, 0).add_paragraph(get_proposal_defense_form[4].panel_full_name)
            panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if get_proposal_defense_form[4].panel_signature_attach == True:
                if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_proposal_defense_form[4].panel_username) + ".png"):
                    if get_proposal_defense_form[4].final_defense_response == "Accepted with Revision":
                        panel_table.cell(5, 1).text = ""
                        panel_table.cell(5, 2).text = ""
                        panel_table.cell(5, 3).text = ""
                        panel_sign = panel_table.cell(5, 1).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[4].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                    
                    elif get_proposal_defense_form[4].final_defense_response == "Deferred with Revision":
                        panel_table.cell(5, 1).text = ""
                        panel_table.cell(5, 2).text = ""
                        panel_table.cell(5, 3).text = ""
                        panel_sign = panel_table.cell(5, 2).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[4].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))

                    elif get_proposal_defense_form[4].final_defense_response == "Not Accepted":
                        panel_table.cell(5, 1).text = ""
                        panel_table.cell(5, 2).text = ""
                        panel_table.cell(5, 3).text = ""
                        panel_sign = panel_table.cell(5, 3).add_paragraph()
                        panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        panel_sign_run = panel_sign.add_run()
                        panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_proposal_defense_form[4].panel_username +'.png',width=Inches(0.708661), height=Inches(0.23622))
                        

                else:
                    context = {
                        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                        "currently_loggedin_user_account": currently_loggedin_user_account,
                        "student_leader_data": get_student_leader_data,
                        "student_leader_full_name": student_leader_full_name,
                        "student_group_members": get_student_group_members,
                        "student_research_title": get_accepted_proposal_title,
                        "proposal_defense_form": get_proposal_defense_form,
                        "response": "sweet faculty member no signature"
                    }

                    return render(request, "student-bet3-critique-form.html", context)
            else:
                if get_proposal_defense_form[4].final_defense_response == "Accepted with Revision":
                    panel_table.cell(5, 2).text = ""
                    panel_table.cell(5, 3).text = ""
                    
                elif get_proposal_defense_form[4].final_defense_response == "Deferred with Revision":
                    panel_table.cell(5, 1).text = ""
                    panel_table.cell(5, 3).text = "" 

                elif get_proposal_defense_form[4].final_defense_response == "Not Accepted":
                    panel_table.cell(5, 1).text = ""
                    panel_table.cell(5, 2).text = ""  
    
    except:
        panel_table.cell(5, 0).text = ""
        panel_table.cell(5, 1).text = ""
        panel_table.cell(5, 2).text = ""
        panel_table.cell(5, 3).text = ""
    

    # Panel 1
    if get_panel_chair_proposal_defense_form[0]:
        panelchair_table.cell(1, 0).text = ""
        panel_name = panelchair_table.cell(1, 0).add_paragraph(get_panel_chair_proposal_defense_form[0].panel_full_name)
        panel_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if get_panel_chair_proposal_defense_form[0].panel_chairman_signature_attach == True:
            if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_panel_chair_proposal_defense_form[0].panel_username) + ".png"):
                panelchair_table.cell(0, 0).text = ""
                panel_sign = panelchair_table.cell(0, 0).add_paragraph()
                panel_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                panel_sign_run = panel_sign.add_run()
                panel_sign_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_panel_chair_proposal_defense_form[0].panel_username +'.png',width=Inches(1.2), height=Inches(0.45))
                
            else:
                context = {
                    "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                    "currently_loggedin_user_account": currently_loggedin_user_account,
                    "student_leader_data": get_student_leader_data,
                    "student_leader_full_name": student_leader_full_name,
                    "student_group_members": get_student_group_members,
                    "student_research_title": get_accepted_proposal_title,
                    "proposal_defense_form": get_proposal_defense_form,
                    "response": "sweet faculty member no signature"
                }

                return render(request, "student-bet3-critique-form.html", context)

    # Create - QR Code
    auth_qr_code = qrcode.make('Research Final Defense Form\nDate:' + get_student_leader_data.research_proposal_defense_date)
    type(auth_qr_code)  
    auth_qr_code.save(current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM-QR.png")

    # INSERT IMAGE
    qr_code = qrcodebox_table.cell(0, 0).add_paragraph()
    qr_code_run = qr_code.add_run()
    qr_code_run.add_picture(current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM-QR.png", width=Inches(1), height=Inches(1))

    doc.save(current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM.docx")
    convert(current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM.pdf")

    # doc.save('/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+'BET5-RESEARCH-FINAL-DEFENSE-FORM.docx')
    # subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', "/home/johnanthonybataller/tupc-research-defense-form-django/static/"+current_user.username+'-BET5-RESEARCH-FINAL-DEFENSE-FORM.docx', "--outdir" ,"/home/johnanthonybataller/tupc-research-defense-form-django/static/"])
    # download_link = "www.ditresearchdefense.online/static/" +current_user.username+'-BET5-RESEARCH-FINAL-DEFENSE-FORM.pdf'

    # filePath =  FilePath(
    #     student_leader_username = current_user.username,
    #     file_path = '/home/johnanthonybataller/tupc-research-defense-form-django/static/'+current_user.username +"-"+panel_username+"-"+panel_response+'-BET5-RESEARCH-FINAL-DEFENSE-FORM.pdf'
    # )
    # filePath.save()

    qr_code_path = current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM-QR.png"

    if os.path.isfile(qr_code_path):
        os.remove(qr_code_path)
        print("Research Final Defense QR - Deleted")
    else:
        print("Research Final Defense QR - Not Found")

    delete_critique_form = current_user.username + "-BET5-RESEARCH-FINAL-DEFENSE-FORM.docx"
    #delete_critique_form = ("/home/johnanthonybataller/tupc-research-defense-form-django/static/" + current_user.username +"-"+panel_username+"-"+panel_response+'-BET3-PROPOSAL-DEFENSE-PANEL-INVITATION.docx')

    if os.path.isfile(delete_critique_form):
        os.remove(delete_critique_form)
        print("Research Final Defense Form - Deleted")
    else:
        print("Research Final Defense Form - Not Found")

    context = {
        "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
        "currently_loggedin_user_account": currently_loggedin_user_account,
        "student_leader_data": get_student_leader_data,
        "student_leader_full_name": student_leader_full_name,
        "student_group_members": get_student_group_members,
        "student_research_title": get_accepted_proposal_title,
        "proposal_defense_form": get_proposal_defense_form,
        "response": "sweet downloaded"
    }

    return render(request, "student-research-final-defense.html", context)


# Student - BET3 - Proposal Defense - Panel Invitation - Download Process
@login_required(login_url="login")
@user_passes_test(lambda u: u.is_student, login_url="login")
def studentAcknowledgmentReceiptDownload(request):
    current_user = request.user
    current_password = current_user.password

    ############## TOPBAR ##############
    topbar_data = topbarProcess(request)
    currently_loggedin_user_full_name = topbar_data[0]
    currently_loggedin_user_account = topbar_data[1]
    ############## TOPBAR ##############

    try:
        get_student_leader_data = StudentLeader.objects.get(username=current_user.username)
        get_group_members = StudentGroupMember.objects.all().filter(student_leader_username=current_user.username)
        get_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
    except:
        return redirect('student-dashboard')

    student_member_list = [get_receipt.student_leader_full_name]

    if get_group_members:
        for group_member in get_group_members:
            student_member_list.append(group_member.student_member_full_name)

    student_member_list.sort()

    print(student_member_list)

    course = get_student_leader_data.course_major_abbr


    doc = Document("static/forms/9-ACKNOWLEDGMENT-RECEIPT.docx")

    title_table = doc.tables[1]
    student_table = doc.tables[2]
    department_box = doc.tables[3]

    title_table.cell(0, 2).text = get_receipt.research_title

    try:
        student_table.cell(0, 2).text = student_member_list[0]
        student_table.cell(0, 4).text = course
    except:
        pass

    try:
        student_table.cell(1, 2).text = student_member_list[1]
        student_table.cell(1, 4).text = course
    except:
        student_table.cell(1, 2).text = ""
        student_table.cell(1, 4).text = ""

    try:
        student_table.cell(2, 2).text = student_member_list[2]
        student_table.cell(2, 4).text = course
    except:
        student_table.cell(2, 2).text = ""
        student_table.cell(2, 4).text = ""

    try:
        student_table.cell(3, 2).text = student_member_list[3]
        student_table.cell(3, 4).text = course
    except:
        student_table.cell(3, 2).text = ""
        student_table.cell(3, 4).text = ""

    try:
        student_table.cell(4, 2).text = student_member_list[4]
        student_table.cell(4, 4).text = course
    except:
        student_table.cell(4, 2).text = ""
        student_table.cell(4, 4).text = ""

    # department_box.cell(1, 2).text = ""
    # dit_name = department_box.cell(1, 2).add_paragraph(get_receipt.dit_head_full_name)
    # dit_name.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # department_box.cell(1, 6).text = ""
    # dit_date = department_box.cell(1, 6).add_paragraph(get_receipt.dit_head_response_date)
    # dit_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    department_box.cell(1, 2).text = get_receipt.dit_head_full_name
    department_box.cell(1, 6).text = get_receipt.dit_head_response_date

    department_box.cell(2, 2).text = get_receipt.adaa_full_name
    department_box.cell(2, 6).text = get_receipt.adaa_response_date

    department_box.cell(3, 2).text = get_receipt.library_full_name
    department_box.cell(3, 6).text = get_receipt.library_response_date
    
    department_box.cell(4, 2).text = get_receipt.research_ext_full_name
    department_box.cell(4, 6).text = get_receipt.research_ext_response_date

    department_box.cell(5, 2).text = get_receipt.adviser_full_name
    department_box.cell(5, 6).text = get_receipt.adviser_response_date

    department_box.cell(6, 2).text = get_receipt.subject_teacher_full_name
    department_box.cell(6, 6).text = get_receipt.subject_teacher_response_date

    if get_receipt.dit_head_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_receipt.dit_head_username) + ".png"):
            department_box.cell(1, 4).text = ''
            panel_signature = department_box.cell(1, 4).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_receipt.dit_head_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            try:
                get_acknowledgement_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
            except:
                pass

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "acknowledgement_receipt": get_acknowledgement_receipt,
                "response": "sweet faculty member no signature"
            }

            return render(request, "student-acknowledgement-receipt-dashboard.html", context)
    
    if get_receipt.adaa_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_receipt.adaa_username) + ".png"):
            department_box.cell(2, 4).text = ''
            panel_signature = department_box.cell(2, 4).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_receipt.adaa_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            try:
                get_acknowledgement_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
            except:
                pass

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "acknowledgement_receipt": get_acknowledgement_receipt,
                "response": "sweet faculty member no signature"
            }

            return render(request, "student-acknowledgement-receipt-dashboard.html", context)
    
    if get_receipt.library_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_receipt.library_username) + ".png"):
            department_box.cell(3, 4).text = ''
            panel_signature = department_box.cell(3, 4).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_receipt.library_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            try:
                get_acknowledgement_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
            except:
                pass

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "acknowledgement_receipt": get_acknowledgement_receipt,
                "response": "sweet faculty member no signature"
            }

            return render(request, "student-acknowledgement-receipt-dashboard.html", context)

    if get_receipt.research_ext_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_receipt.research_ext_username) + ".png"):
            department_box.cell(4, 4).text = ''
            panel_signature = department_box.cell(4, 4).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_receipt.research_ext_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            try:
                get_acknowledgement_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
            except:
                pass

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "acknowledgement_receipt": get_acknowledgement_receipt,
                "response": "sweet faculty member no signature"
            }

            return render(request, "student-acknowledgement-receipt-dashboard.html", context)
    
    if get_receipt.adviser_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_receipt.adviser_username) + ".png"):
            department_box.cell(5, 4).text = ''
            panel_signature = department_box.cell(5, 4).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_receipt.adviser_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            try:
                get_acknowledgement_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
            except:
                pass

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "acknowledgement_receipt": get_acknowledgement_receipt,
                "response": "sweet faculty member no signature"
            }

            return render(request, "student-acknowledgement-receipt-dashboard.html", context)
    
    if get_receipt.subject_teacher_signature == True:
        # Check if Panel E-Sign Exist
        if os.path.exists("uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/" + str(get_receipt.subject_teacher_username) + ".png"):
            department_box.cell(6, 4).text = ''
            panel_signature = department_box.cell(6, 4).add_paragraph()
            panel_signature_run = panel_signature.add_run()
            panel_signature_run.add_picture('uhsG1tCRrm3fUHcG4dyEMDDq31WQULMNJkSGQFq0oiV5vvhui9/'+get_receipt.subject_teacher_username +'.png',width=Inches(1.2), height=Inches(0.45))
        else:
            try:
                get_acknowledgement_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
            except:
                pass

            context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "acknowledgement_receipt": get_acknowledgement_receipt,
                "response": "sweet faculty member no signature"
            }

            return render(request, "student-acknowledgement-receipt-dashboard.html", context)

    doc.save(current_user.username + "-ACKNOWLEDGEMENT-RECEIPT.docx")
    convert(current_user.username + "-ACKNOWLEDGEMENT-RECEIPT.docx")

    filePath = FilePath(
        student_leader_username=current_user.username, 
        file_path=current_user.username + "-ACKNOWLEDGEMENT-RECEIPT.pdf"
        )
    filePath.save()

    os.startfile(current_user.username + "-ACKNOWLEDGEMENT-RECEIPT.pdf")

    try:
        get_acknowledgement_receipt = AcknowledgementReceipt.objects.get(student_leader_username=current_user.username)
    except:
        pass

    context = {
                "currently_loggedin_user_full_name": currently_loggedin_user_full_name,
                "currently_loggedin_user_account": currently_loggedin_user_account,
                "student_leader_data": get_student_leader_data,
                "acknowledgement_receipt": get_acknowledgement_receipt,
                "response": "sweet downloaded"
            }

    return render(request, "student-acknowledgement-receipt-dashboard.html", context)


@login_required(login_url="login")
def topbarProcess(request):

    currently_loggedin_user = request.user
    currently_loggedin_user_middle_name = currently_loggedin_user.middle_name
    currently_loggedin_user_middle_initial = None

    currently_loggedin_user_full_name = None
    currently_loggedin_user_account = None

    if currently_loggedin_user_middle_name == "":
        currently_loggedin_user_full_name = currently_loggedin_user.first_name + " " + currently_loggedin_user.last_name + " " + currently_loggedin_user.suffix

    else:
        currently_loggedin_user_middle_initial = currently_loggedin_user_middle_name[0]
        currently_loggedin_user_full_name = currently_loggedin_user.first_name + " " + currently_loggedin_user_middle_initial + ". " + currently_loggedin_user.last_name + " " + currently_loggedin_user.suffix

    if currently_loggedin_user.is_student == 1:
        currently_loggedin_user_account = "Student"

    elif currently_loggedin_user.is_administrator == 1:
        currently_loggedin_user_account = "Administrator"

    elif currently_loggedin_user.is_academic_affairs == 1:
        currently_loggedin_user_account = "Academic Affairs"

    elif currently_loggedin_user.is_library == 1:
        currently_loggedin_user_account = "Library"

    elif currently_loggedin_user.is_research_extension == 1:
        currently_loggedin_user_account = "Research & Extension"

    return (currently_loggedin_user_full_name, currently_loggedin_user_account)


@login_required(login_url="login")
def fullNameProcess(request, id):

    print(id)

    faculty_full_name = ""

    try:
        get_faculty_data = User.objects.get(username=id)
        print("pass")

    except:
        return faculty_full_name

    if get_faculty_data.middle_name == "":
        faculty_full_name = get_faculty_data.honorific + " " + get_faculty_data.first_name + " " + get_faculty_data.last_name + " " + get_faculty_data.suffix
        return faculty_full_name
    else:
        faculty_full_name = get_faculty_data.honorific + " " + get_faculty_data.first_name + " " + get_faculty_data.middle_name[0] + ". " + get_faculty_data.last_name + " " + get_faculty_data.suffix
        return faculty_full_name